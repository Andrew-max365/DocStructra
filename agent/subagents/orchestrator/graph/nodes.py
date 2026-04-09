# agent/graph/nodes.py
from __future__ import annotations

from typing import Any, Dict, List, Set
from agent.graph.react_schemas import Action, ActionPlan, Observation, GraphState


def ingest_node(state: GraphState) -> dict:
    from core.parser import parse_docx_to_blocks
    from core.judge import rule_based_labels
    from core.docx_utils import iter_all_paragraphs
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    import logging

    logger = logging.getLogger(__name__)

    doc, blocks = parse_docx_to_blocks(state["input_path"])
    # 极速第一步：打上纯规则标签
    labels = rule_based_labels(blocks, doc=doc)
    labels["_source"] = "unified_workflow"

    # ── LLM 正文范围识别：识别正文起止，并插入 {body} 标签 ──
    try:
        from agent.llm_client import LLMClient
        
        all_paragraphs = [p.text for p in iter_all_paragraphs(doc)]
        client = LLMClient()
        range_res = client.call_body_range_identification(all_paragraphs)
        
        start_idx = range_res.get("start_index", 0)
        end_idx = range_res.get("end_index", len(all_paragraphs) - 1)
        
        # 为了不破坏后续 block 的映射，我们在 doc 中插入标记段落。
        # 注意：从后往前插入，避免前面的索引变动。
        
        # 插入 {/body}
        # 如果 end_idx 是最后一段，则在最后追加；否则在 end_idx+1 段之前插
        paragraphs_list = list(iter_all_paragraphs(doc))
        if end_idx < len(paragraphs_list):
            target_p = paragraphs_list[end_idx]
            new_p = OxmlElement('w:p')
            target_p._p.addnext(new_p)
            new_para = Paragraph(new_p, target_p._parent)
            new_para.add_run("{/body}")
        
        # 插入 {body}
        if start_idx < len(paragraphs_list):
            target_p = paragraphs_list[start_idx]
            new_p = OxmlElement('w:p')
            target_p._p.addprevious(new_p)
            new_para = Paragraph(new_p, target_p._parent)
            new_para.add_run("{body}")
            
        logger.info(f"[body_range] 识别到正文范围: {start_idx}-{end_idx}, 已插入标记")
        
        # 重新解析 blocks，因为我们改变了文档结构（增加了两个段落）
        # 重新调用 parse_docx_to_blocks 有点重，但最保险
        # 这里的 blocks 用于后续的 label 映射，需要与 doc 同步
        _, blocks = parse_docx_to_blocks(state["input_path"], doc=doc)
        # 重新生成 labels
        labels = rule_based_labels(blocks, doc=doc)
        labels["_source"] = "unified_workflow_with_body_tags"

    except Exception as e:
        logger.warning(f"[body_range] 识别正文范围失败: {e}")

    return {
        "doc": doc,
        "blocks": blocks,
        "labels": labels,
    }


def trigger_node(state: GraphState) -> dict:
    """雷达扫描：检查规则排版后是否存在疑似异常，决定是否唤醒大模型"""
    blocks = state["blocks"]
    rule_labels = state["labels"]

    triggered_indices: Set[int] = set()
    reasons: List[str] = []

    # 触发条件1：规则无法判定的 unknown
    unknown_blocks = [b for b in blocks if rule_labels.get(b.block_id) == "unknown"]
    if unknown_blocks:
        triggered_indices.update(b.paragraph_index for b in unknown_blocks)
        reasons.append(f"存在 {len(unknown_blocks)} 个未知格式段落")

    # 触发条件2：疑似被误分类的长标题 (>30字)
    ambiguous_headings = [b for b in blocks if rule_labels.get(b.block_id) in ("h2", "h3") and len(b.text or "") > 30]
    if ambiguous_headings:
        triggered_indices.update(b.paragraph_index for b in ambiguous_headings)
        reasons.append(f"存在 {len(ambiguous_headings)} 个超长标题")

    # 触发条件3：连续短句 (可能是未识别列表)
    run = []
    sorted_blocks = sorted(blocks, key=lambda x: x.paragraph_index)
    for b in sorted_blocks:
        if rule_labels.get(b.block_id) == "body" and 0 < len((b.text or "").strip()) <= 60:
            run.append(b)
        else:
            if len(run) >= 3:
                triggered_indices.update(rb.paragraph_index for rb in run)
                reasons.append(f"发现连续短正文(段落{run[0].paragraph_index}~{run[-1].paragraph_index})，可能需要结构化")
            run = []
    if len(run) >= 3:
        triggered_indices.update(rb.paragraph_index for rb in run)
        reasons.append(f"发现连续短正文(段落{run[0].paragraph_index}~{run[-1].paragraph_index})，可能需要结构化")

    # 触发条件4：空段落密度过高 (>30%)
    empty_blocks = [b for b in blocks if not (b.text or "").strip()]
    if blocks and len(empty_blocks) / len(blocks) > 0.3:
        reasons.append(f"空段落占比过高({len(empty_blocks)}/{len(blocks)})，建议清理排版")

    # 触发条件5：字号极其不统一 (异常波动)
    # 若同一角色的段落存在 >3 种不同的字号，可能是直接复制粘贴导致格式紊乱
    role_sizes = {}
    from core.docx_utils import iter_all_paragraphs
    doc_paragraphs = list(iter_all_paragraphs(state["doc"]))
    
    for b in blocks:
        role = rule_labels.get(b.block_id, "body")
        try:
            p = doc_paragraphs[b.paragraph_index]
            if p.runs and p.runs[0].font.size:
                size_pt = p.runs[0].font.size.pt
                if size_pt:
                    if role not in role_sizes:
                        role_sizes[role] = set()
                    role_sizes[role].add(size_pt)
        except IndexError:
            pass
    
    for role, sizes in role_sizes.items():
        if len(sizes) > 3:
            reasons.append(f"角色 '{role}' 存在多种不同字号 {sizes}，建议统一样式")
            # 不全量 trigger，仅做 warning 记录，供 reason_node 参考


    needs_llm = len(triggered_indices) > 0

    hybrid_triggers = {
        "triggered": needs_llm,
        "reasons": reasons,
        "triggered_paragraph_count": len(triggered_indices),
    }

    # 注意：这里我们提前初始化校对问题列表，防止后续没调用 LLM 时报错
    return {
        "needs_llm": needs_llm,
        "triggered_indices": sorted(list(triggered_indices)),
        "hybrid_triggers": hybrid_triggers,
        "proofread_issues": state.get("proofread_issues", [])
    }


def reason_node(state: GraphState) -> dict:
    """大模型思考节点：仅处理被雷达标记的段落，并收集文字校对建议"""
    from agent.llm_client import LLMClient
    from core.judge import SmartJudge
    from core.docx_utils import iter_all_paragraphs
    import traceback

    iteration = state["current_iter"] + 1
    blocks = state["blocks"]
    doc = state["doc"]
    rule_labels = state["labels"]
    triggered_indices = state.get("triggered_indices", [])
    proofread_issues = state.get("proofread_issues", [])
    if not triggered_indices:
        return {
            "thoughts": list(state.get("thoughts", [])) + [f"迭代 {iteration}: 未命中异常段落，跳过大模型调用。"],
            "actions": list(state.get("actions", [])) + [[{"action_type": "no_op", "block_id": -1}]],
            "proofread_issues": proofread_issues,
            "current_iter": iteration,
        }

    actions_to_take = []
    thought = f"迭代 {iteration}: 雷达标记了 {len(triggered_indices)} 个异常段落，唤醒大模型进行深层结构分析与文本校对。"

    # 将视觉反馈注入 LLM 的分析上下文
    visual_feedback = state.get("visual_feedback_for_reason")
    if visual_feedback:
        thought += f" [视觉审查反馈：{visual_feedback}]"

    try:
        client = LLMClient()
        all_paragraphs = [p.text for p in iter_all_paragraphs(doc)]
        smart_judge = SmartJudge()

        # 1. 结构分析（纠正排版标签）
        struct_res = client.call_structure_analysis(all_paragraphs, triggered_indices)
        llm_by_index = {pr.paragraph_index: {"role": pr.role, "confidence": pr.confidence} for pr in
                        struct_res.paragraphs}

        index_to_block = {b.paragraph_index: b for b in blocks}
        for pidx in triggered_indices:
            b = index_to_block.get(pidx)
            if b:
                orig_role = rule_labels.get(b.block_id, "body")
                llm_dict = llm_by_index.get(pidx, {})
                if llm_dict:
                    final_role = smart_judge.arbitrate(b.text or "", orig_role, llm_dict)
                    if final_role != orig_role:
                        actions_to_take.append({
                            "action_type": "set_role",
                            "block_id": b.block_id,
                            "params": {"role": final_role}
                        })

        # 2. 顺便进行文字校对（收集错别字建议交给用户）
        proof_res = client.call_proofread(all_paragraphs, triggered_indices)
        proofread_issues = [issue.model_dump() for issue in proof_res.issues]

        if not actions_to_take:
            actions_to_take.append({"action_type": "no_op", "block_id": -1})
            thought += " 大模型认为现有结构合理，无需修改。"

    except Exception as e:
        thought += f" 调用大模型失败: {e}\n{traceback.format_exc()}"
        actions_to_take.append({"action_type": "no_op", "block_id": -1})

    return {
        "thoughts": list(state.get("thoughts", [])) + [thought],
        "actions": list(state.get("actions", [])) + [actions_to_take],
        "proofread_issues": proofread_issues,
        "current_iter": iteration,
    }


def act_node(state: GraphState) -> dict:
    """执行节点：将大模型的判定覆盖回 labels"""
    labels = dict(state["labels"])
    latest_actions = state["actions"][-1] if state.get("actions") else []

    for action in latest_actions:
        if action.get("action_type") == "set_role":
            block_id = action.get("block_id", -1)
            role = action.get("params", {}).get("role")
            if block_id >= 0 and role:
                labels[block_id] = role
        elif action.get("action_type") == "fix_heading_level":
            block_id = action.get("block_id", -1)
            level = action.get("params", {}).get("level")
            if block_id >= 0 and level:
                labels[block_id] = f"h{level}"

    return {"labels": labels}


def validate_node(state: GraphState) -> dict:
    from core.formatter import apply_formatting
    from core.writer import save_docx
    from core.spec import load_spec

    doc = state["doc"]
    blocks = state["blocks"]
    labels = state["labels"]
    spec = load_spec(state["spec_path"], overrides=state.get("overrides"))

    errors = []
    passed = True
    report = state.get("report", {})

    try:
        report = apply_formatting(doc, blocks, labels, spec)
        if not report:
            errors.append("apply_formatting 返回空报告")
            passed = False
        else:
            # ✨ 关键：将收集到的 LLM 数据注入报告，让前端能渲染按钮！
            report["hybrid_triggers"] = state.get("hybrid_triggers", {})
            if state.get("proofread_issues"):
                report["llm_proofread"] = {"issues": state["proofread_issues"]}
    except Exception as e:
        import traceback
        errors.append(f"格式化失败: {e}\n{traceback.format_exc()}")
        passed = False

    try:
        save_docx(doc, state["output_path"])
    except Exception as e:
        errors.append(f"保存文档失败: {e}")
        passed = False

    obs = Observation(
        iteration=state["current_iter"],
        passed=passed,
        errors=errors,
        summary=f"排版验证结束: {'成功' if passed else '失败'}",
    )

    return {
        "observations": list(state.get("observations", [])) + [obs.model_dump()],
        "errors": errors,
        "passed": passed,
        "report": report,
        # finished 现在由 reflect_router 决定（如果不触发 reflect 则看这里）
        "finished": (not state.get("visual_review_enabled", False)) and (passed or state.get("current_iter", 0) >= state.get("max_iters", 0)),
    }


def reflect_node(state: GraphState) -> dict:
    """视觉反思节点：强制使用具备 Vision 能力的 API 进行审查，同时保留所有历史记录"""

    # 🚀 视觉审查未启用时，跳过整个节点逻辑，直接返回安全的默认状态
    if not state.get("visual_review_enabled", False):
        return {
            "visual_review_result": None,
            "reflection_history": list(state.get("reflection_history", [])),
            "reflection_count": state.get("reflection_count", 0),
            "thoughts": list(state.get("thoughts", [])) + ["视觉审查未启用，跳过反思节点。"],
            "actions": list(state.get("actions", [])) + [[{"action_type": "no_op", "block_id": -1}]],
            "report": dict(state.get("report", {})),
            "visual_feedback_for_reason": None,
        }

    import traceback
    import os
    import asyncio
    import tempfile
    import shutil
    from agent.llm_client import LLMClient
    from core.spec import load_spec
    from agent.spec_summarizer import summarize_spec
    from config import REFLECTION_MAX_ITERS
    from agent.visual_reviewer import docx_to_pdf, pdf_to_images

    # 1. --- 完美保留你原有的状态获取 ---
    output_path = state["output_path"]
    reflection_count = state.get("reflection_count", 0)
    current_iter = state.get("current_iter", 0)
    history = list(state.get("reflection_history", []))
    report = dict(state.get("report", {}))

    thought = f"反思迭代 {reflection_count}: 启动视觉模型(Gemini/GPT)进行多模态审计。"
    actions_to_take = []
    vr_res_dict = None

    try:
        # 2. --- 准备审计上下文 ---
        spec = load_spec(state["spec_path"], overrides=state.get("overrides"))
        spec_summary = summarize_spec(spec)

        # 3. --- 将 DOCX 转为图片 (先转 PDF，再转图片，只转一页) ---
        tmp_dir = tempfile.mkdtemp(prefix="vision_audit_")

        try:
            pdf_path = docx_to_pdf(output_path, output_dir=tmp_dir)
            img_paths = pdf_to_images(pdf_path, max_pages=1, output_dir=tmp_dir)

            if not img_paths:
                raise ValueError("未能生成审计图片，请检查 LibreOffice 和 poppler 是否配置正确")

            # 4. 🌟 --- 核心替换：调用视觉模型“眼睛” ---
            # 实例化 LLMClient 后用 asyncio.run 运行异步方法
            client = LLMClient()
            vision_feedback = asyncio.run(client.call_vision_audit(img_paths[0], spec_summary))

            # 如果你的 call_vision_audit 是写成了静态方法（@staticmethod），请注释掉上面两行，用下面这行：
            # vision_feedback = asyncio.run(LLMClient.call_vision_audit(img_paths[0], spec_summary))

            # 5. --- 解析视觉反馈 (将其模拟为你原来的 vr_res 结构) ---
            needs_reformat = any(word in vision_feedback for word in ["错误", "不符合", "修改", "重新", "不一致"])

            vr_res_dict = {
                "overall_score": 8.5 if not needs_reformat else 6.0,
                "summary": vision_feedback,
                "needs_reformat": needs_reformat,
                "issues": []
            }

            thought += f" 视觉评分: {vr_res_dict['overall_score']:.1f}。反馈：{vision_feedback[:50]}..."

            # 将结果注入报告，方便前端渲染
            report["visual_review"] = vr_res_dict

            # 6. --- 完美保留你原有的重排动作生成逻辑 ---
            if needs_reformat and reflection_count < REFLECTION_MAX_ITERS:
                thought += " 视觉审计未通过，生成修正指令传回给大脑(DeepSeek)..."
                actions_to_take.append({
                    "action_type": "adjust_layout_by_vision",
                    "block_id": -1,
                    "params": {"feedback": vision_feedback}
                })
            else:
                thought += " 视觉审计通过或已达最大反思次数，流程结束。"

        finally:
            # 无论成功失败，最后清理掉这个临时文件夹里的 PDF 和 图片，防止撑爆硬盘
            shutil.rmtree(tmp_dir, ignore_errors=True)

    except Exception as e:
        thought += f" 视觉审计失败: {e}\n{traceback.format_exc()}"
        vr_res_dict = {"needs_reformat": False, "summary": str(e)}

    # 7. --- 完美保留你的历史记录追加逻辑 ---
    history.append({
        "reflection_count": reflection_count,
        "thought": thought,
        "result": vr_res_dict
    })

    # 为下一跳做好准备
    if not actions_to_take:
        actions_to_take.append({"action_type": "no_op", "block_id": -1})

    return {
        "visual_review_result": vr_res_dict,
        "reflection_history": history,
        "reflection_count": reflection_count + 1,
        "thoughts": list(state.get("thoughts", [])) + [thought],
        "actions": list(state.get("actions", [])) + [actions_to_take],
        "report": report,
        "visual_feedback_for_reason": vr_res_dict["summary"] if vr_res_dict else None,
    }

def route_trigger(state: GraphState) -> str:
    """智能路由：按需唤醒 LLM"""
    if state.get("needs_llm"):
        return "reason"  # 去呼叫大模型
    return "validate"  # 极其规整，直接格式化出图！


def retry_router(state: GraphState) -> str:
    """旧的直接验证路由（Visual Review 关闭时使用）"""
    # 🚀 【核心修复】：只要出现了 errors（说明代码异常崩溃），立刻结束，不要让大模型重试！
    if state["passed"] or len(state.get("errors", [])) > 0 or state.get("current_iter", 0) >= state.get("max_iters", 0):
        return "end"
    return "reason"


def reflect_router(state: GraphState) -> str:
    """
    视觉反思路由：决定是结束，还是触发下一轮 Reason 结构修正。
    """
    if not state.get("visual_review_enabled", False):
        return retry_router(state)
        
    vr = state.get("visual_review_result")
    from config import REFLECTION_MAX_ITERS
    reflection_count = state.get("reflection_count", 0)
    
    # 无结果或已满次数 -> 结束
    if not vr or reflection_count >= REFLECTION_MAX_ITERS:
        return "end"
        
    # 如果需要重排，且没有超过次数 -> 重新规划
    needs_reformat = vr.get("needs_reformat", False)
    if needs_reformat:
        return "reason"
        
    return "end"
