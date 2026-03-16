# ui/chainlit_app.py
"""
Chainlit 前端入口 — MyAgent ReAct 文档格式化（增强版）。

功能：
  1. 点击按钮选择排版模式（无需手动输入命令）。
  2. 直接上传 .docx 文件即可处理，无需同时输入文字。
  3. Diff 视图：直接在页面中渲染修改前后对比（GFM ~~删除线~~ → 建议，含段落上下文）。
  4. 通用聊天：不上传文件时，支持直接与 LLM 对话（流式输出）。

启动方式：
    chainlit run ui/chainlit_app.py
"""
from __future__ import annotations

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(__file__)))

import asyncio
import copy
import json
import re
import tempfile
from typing import Any, Dict, List, Set

try:
    import chainlit as cl
except ImportError as e:
    raise ImportError("chainlit 未安装，请运行: pip install chainlit") from e

from config import LLM_MODE, REACT_MAX_ITERS, LLM_API_KEY, LLM_BASE_URL, LLM_MODEL
from service.format_service import format_docx_bytes
from ui.diff_utils import (
    build_diff_items,
    parse_rejected_numbers,
    apply_and_save_proofread,
    generate_structural_diff,
    _ACCEPT_ALL_PATTERNS,
    DiffItem,
)

# Session state keys
_KEY_MAX_ITERS = "max_iters"
_KEY_STATE = "ui_state"        # "ready" | "awaiting_feedback"
_KEY_INPUT_BYTES = "input_bytes"
_KEY_OUTPUT_BYTES = "output_bytes"
_KEY_FILENAME = "filename"
_KEY_ISSUES = "pending_issues"
_KEY_DIFF_ITEMS = "diff_items"
_KEY_REPORT = "pending_report"
_KEY_CHAT_HISTORY = "chat_history"
_KEY_SPEC_OVERRIDES = "spec_overrides"
_KEY_SPEC_PATH = "spec_path"
_KEY_HFT_ACTIONS = "hft_actions"  # 待应用的页眉/页脚/页码/目录操作（无文档时暂存）

# Maximum number of chat messages (user+assistant turns) to keep in session context.
_MAX_CHAT_HISTORY = 20


def _deep_merge_dicts(base: Dict[str, Any], update: Dict[str, Any]) -> Dict[str, Any]:
    """深度合并两个配置字典，update 中的叶子值覆盖 base，不整层替换。委托给 core.spec._deep_merge。"""
    from core.spec import _deep_merge
    return _deep_merge(base, update)


@cl.on_chat_start
async def on_chat_start():
    cl.user_session.set(_KEY_MAX_ITERS, REACT_MAX_ITERS)
    cl.user_session.set(_KEY_STATE, "ready")
    cl.user_session.set(_KEY_CHAT_HISTORY, [])
    cl.user_session.set(_KEY_SPEC_OVERRIDES, {})
    cl.user_session.set(_KEY_SPEC_PATH, "specs/default.yaml")
    cl.user_session.set(_KEY_HFT_ACTIONS, {})

    await cl.Message(
        content=(
            "👋 欢迎使用 **Structura 智能文档排版助手**！\n\n"
            "直接上传 `.docx` 文件即可开始全自动排版（极速规则 + 大模型智能纠错）。\n\n"
            "📝 **排版指令** `/f`：使用 `/f` 或 `/format` 开头，让 AI 识别您的自然语言排版需求并应用，例如：\n"
            "> `/f 把大标题改成红色，正文字号改成 14，在页眉中写上华中科技大学，并添加页码`\n\n"
            "🔍 **审阅指令** `/r`：使用 `/r` 或 `/review` 开头，审阅文档排版问题；附带要求时只做增量修改，例如：\n"
            "> `/r` — 仅审阅，列出排版问题\n"
            "> `/r 把正文行距改为 1.5 倍` — 审阅 + 只改行距，不动其他格式\n\n"
            "💬 **自由聊天**：直接发送消息（不加前缀）与我对话。"
        )
    ).send()



@cl.action_callback("accept_all_action")
async def on_accept_all(action: cl.Action):
    if cl.user_session.get(_KEY_STATE) != "awaiting_feedback":
        return
    await _execute_feedback("accept_all", [])

@cl.action_callback("reject_all_action")
async def on_reject_all(action: cl.Action):
    if cl.user_session.get(_KEY_STATE) != "awaiting_feedback":
        return
    await _execute_feedback("reject_all", [])


@cl.on_message
async def on_message(message: cl.Message):
    state = cl.user_session.get(_KEY_STATE, "ready")

    # Allow uploading a new file even while awaiting feedback (starts fresh)
    docx_file = None
    for f in (message.elements or []):
        if hasattr(f, "name") and f.name.lower().endswith(".docx"):
            docx_file = f
            break

    if state == "awaiting_feedback" and docx_file is None:
        await _handle_feedback(message)
        return

    text = message.content.strip()

    # ── File upload (text is optional) ──────────────────────────────────────
    if docx_file is not None:
        max_iters = cl.user_session.get(_KEY_MAX_ITERS, REACT_MAX_ITERS)
        overrides = cl.user_session.get(_KEY_SPEC_OVERRIDES, {})
        spec_path = cl.user_session.get(_KEY_SPEC_PATH, "specs/default.yaml")

        with open(docx_file.path, "rb") as fp:
            input_bytes = fp.read()

        cl.user_session.set(_KEY_INPUT_BYTES, input_bytes)
        cl.user_session.set(_KEY_FILENAME, docx_file.name)
        # 新文件上传时清空上次的输出，避免增量操作误用旧文件
        cl.user_session.set(_KEY_OUTPUT_BYTES, None)

        # ── 若用户在发送文件时附带了文字要求，根据前缀分发 ──
        if text:
            # /f 前缀 → 排版指令（解析全量 spec + HFT 需求后处理文档）
            if _is_format_command(text):
                cmd_content = _extract_format_content(text)
                if cmd_content:
                    thinking_msg = cl.Message(content="⏳ 正在解析您的排版要求...")
                    await thinking_msg.send()

                    try:
                        from agent.intent_parser import parse_formatting_request
                        parsed = await parse_formatting_request(cmd_content, current_spec_path=spec_path)
                        formatting_intent = parsed.get("overrides", {})
                        hft_actions = parsed.get("hft_actions", {})
                        routed_spec = parsed.get("spec_path", spec_path)
                    except Exception as e:
                        formatting_intent = None
                        hft_actions = {}
                        routed_spec = spec_path
                        print(f"解析排版意图异常: {e}")

                    if formatting_intent:
                        new_overrides = _deep_merge_dicts(copy.deepcopy(overrides), formatting_intent)
                        cl.user_session.set(_KEY_SPEC_OVERRIDES, new_overrides)
                        cl.user_session.set(_KEY_SPEC_PATH, routed_spec)
                        overrides = new_overrides
                        spec_path = routed_spec

                        pretty = json.dumps(formatting_intent, ensure_ascii=False, indent=2)
                        thinking_msg.content = (
                            f"✅ **排版指令已确认！**\n\n"
                            f"```json\n{pretty}\n```\n\n"
                            f"📚 模板：`{routed_spec}`\n"
                            f"⏳ 正在处理文档..."
                        )
                        await thinking_msg.update()
                    else:
                        await thinking_msg.remove()

                    await _process_file(input_bytes, docx_file.name, max_iters, overrides=overrides if overrides else None, spec_path=spec_path, _hft_pending=bool(hft_actions))

                    # 如果有 HFT 需求（页眉/页脚/页码/目录），在格式化完成后继续处理
                    if hft_actions:
                        current_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES) or input_bytes
                        await _handle_header_footer_toc(
                            current_bytes, cmd_content, docx_file.name, pre_parsed_cmd=hft_actions
                        )
                        # 如果没有待处理的校对建议（未进入 awaiting_feedback），立即提供下载
                        if cl.user_session.get(_KEY_STATE) != "awaiting_feedback":
                            out_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES)
                            report = cl.user_session.get(_KEY_REPORT, {})
                            await _provide_download(out_bytes, report, docx_file.name, applied=0)
                else:
                    # /f 但没有内容 → 直接用已有偏好处理
                    await _process_file(input_bytes, docx_file.name, max_iters, overrides=overrides if overrides else None, spec_path=spec_path)
                return

            # /r 前缀 → 审阅指令（先审阅，再可选增量修改）
            if _is_review_command(text):
                review_content = _extract_review_content(text)
                await _apply_review_flow(input_bytes, review_content, docx_file.name)
                return

            # 无前缀或其他文字 → 当作普通备注，直接全量处理文档（不解析排版要求）
            # 将文字发给聊天，文档做标准处理
            await _handle_chat(text)

        await _process_file(input_bytes, docx_file.name, max_iters, overrides=overrides if overrides else None, spec_path=spec_path, _hft_pending=bool(cl.user_session.get(_KEY_HFT_ACTIONS, {})))

        # 若有之前暂存的 HFT 操作（无文档时通过 /f 记录），应用后清空
        saved_hft = cl.user_session.get(_KEY_HFT_ACTIONS, {})
        if saved_hft:
            current_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES) or input_bytes
            await _handle_header_footer_toc(current_bytes, "", docx_file.name, pre_parsed_cmd=saved_hft)  # no extra user_text needed
            cl.user_session.set(_KEY_HFT_ACTIONS, {})
            # 如果没有待处理的校对建议，提供下载
            if cl.user_session.get(_KEY_STATE) != "awaiting_feedback":
                out_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES)
                report = cl.user_session.get(_KEY_REPORT, {})
                await _provide_download(out_bytes, report, docx_file.name, applied=0)
        return

    # ── General chat fallback ────────────────────────────────────────────────
    if text:
        await _handle_chat(text)
    else:
        await cl.Message(
            content="💡 请上传 `.docx` 文件开始排版，或直接发送消息与我对话。",
        ).send()


# ── Core processing ────────────────────────────────────────────────────────

async def _process_file(
        input_bytes: bytes,
        filename: str,
        max_iters: int,
        overrides: dict = None,
        spec_path: str = "specs/default.yaml",
        _hft_pending: bool = False,
) -> None:
    """Run the formatting pipeline and display results.

    :param _hft_pending: When True, skip the immediate _provide_download call when no
                         diff cards are shown, so the caller can apply HFT changes first
                         and then provide a single final download.
    """

    processing_msg = cl.Message(content=f"🚀 任务已启动：正在全自动处理文档...")
    await processing_msg.send()

    try:
        # 🌟 调用流式流程
        out_bytes, report = await _run_react_with_steps(
            input_bytes, filename, max_iters, overrides=overrides, spec_path=spec_path
        )
    except Exception as e:
        processing_msg.content = f"❌ 处理失败：{e}"
        await processing_msg.update()
        return

    # 隐藏刚才的过渡消息
    await processing_msg.remove()

    # 1. 保存当前状态，防止用户点按钮时找不到数据
    cl.user_session.set(_KEY_OUTPUT_BYTES, out_bytes)
    cl.user_session.set(_KEY_REPORT, report)
    cl.user_session.set(_KEY_FILENAME, filename)

    try:
        # 2. 生成排版结构变更摘要
        from ui.diff_utils import generate_structural_diff, build_diff_items
        struct_diff = generate_structural_diff(report)
        if struct_diff:
            await cl.Message(
                content=f"### 📐 排版格式化变更摘要\n\n{struct_diff}"
            ).send()

        # 3. 提取错别字校对建议
        # ⚠️ 注意：这里必须用 get 的链式调用，防止大模型没返回 proofread 导致报错
        raw_issues = report.get("llm_proofread", {}).get("issues", [])
        diff_items = build_diff_items(raw_issues)

        cl.user_session.set(_KEY_ISSUES, raw_issues)
        cl.user_session.set(_KEY_DIFF_ITEMS, diff_items)

        # 4. 如果有错别字建议，展示卡片让用户选；如果没有，直接给下载链接！
        if diff_items:
            await _show_diff_cards(diff_items)
            cl.user_session.set(_KEY_STATE, "awaiting_feedback")
        elif not _hft_pending:
            # 没有任何错别字，直接出锅！
            await _provide_download(out_bytes, report, filename, applied=0)
    except Exception as e:
        # 后处理出错时，仍然尝试提供下载，不让用户白等
        await cl.Message(content=f"⚠️ 后处理阶段出现异常：{e}，但文档已成功排版。").send()
        await _provide_download(out_bytes, report, filename, applied=0)


async def _run_react_with_steps(
        input_bytes: bytes,
        filename: str,
        max_iters: int,
        overrides: dict = None,
        spec_path: str = "specs/default.yaml",
) -> tuple:
    """Run the ReAct agent and display progress dynamically via streaming."""
    tmp_in = tmp_out = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(input_bytes)
            tmp_in = f.name
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_out = f.name

        # 🚀 引入我们刚刚写好的流式生成函数
        from agent.graph.workflow import run_react_agent_stream

        result_state = {}

        # 替代原本阻塞式的 asyncio.to_thread，使用 async for 优雅地“倾听”Agent的进展
        async for event in run_react_agent_stream(
                tmp_in, tmp_out, spec_path=spec_path, max_iters=max_iters, overrides=overrides
        ):
            for node_name, state_update in event.items():
                # 记录最终状态
                result_state.update(state_update)

                # 🚀 核心：根据不同节点的完成状态，向前端“直播”
                if node_name == "ingest":
                    await cl.Message(content="✅ **阶段 1/4**：读取成功，已完成基础规则格式解析。").send()
                elif node_name == "trigger":
                    if state_update.get("needs_llm"):
                        await cl.Message(
                            content="🔍 **阶段 2/4**：雷达扫描到异常段落！\n"
                                    "> 正在唤醒大模型进行深度结构分析与错别字校对...\n"
                                    "*(此过程大约需要 30~60 秒，请您先喝口水 ☕)* ⏳"
                        ).send()
                    else:
                        await cl.Message(
                            content="⚡ **阶段 2/4**：文档结构极其清晰！\n"
                                    "> 无需大模型介入，已为您切换至极速排版模式。"
                        ).send()
                elif node_name == "reason":
                    await cl.Message(content="🧠 **阶段 3/4**：大模型结构分析完毕，正在为您应用智能排版策略...").send()
                elif node_name == "validate":
                    passed = state_update.get("passed", False)
                    if passed:
                        await cl.Message(content="✨ **阶段 4/4**：所有排版格式应用成功！").send()
                    else:
                        errors = state_update.get("errors", [])
                        await cl.Message(content=f"⚠️ **排版异常警告**：\n{errors}").send()
                elif node_name == "reflect":
                    await cl.Message(content="👀 **附加阶段**：视觉多模态审查完毕。").send()

        # 读取最终生成的二进制文档
        with open(tmp_out, "rb") as f:
            out_bytes = f.read()

        return out_bytes, result_state.get("report", {})

    except Exception as e:
        import traceback
        await cl.Message(
            content=f"⚠️ 排版流水线发生崩溃: {e}\n```\n{traceback.format_exc()}\n```"
        ).send()

        # 原有的安全回退机制
        from service.format_service import format_docx_bytes
        out_bytes, report = await asyncio.to_thread(
            format_docx_bytes,
            input_bytes, filename_hint=filename, label_mode="hybrid",
            overrides=overrides,
        )
        return out_bytes, report

    finally:
        for p in (tmp_in, tmp_out):
            if p:
                try:
                    os.remove(p)
                except OSError:
                    pass

async def _show_diff_cards(diff_items: List[DiffItem]) -> None:
    """Display diff items as plain markdown with action buttons at the bottom."""
    lines = [f"### 🔍 LLM 校对建议（共 {len(diff_items)} 条）\n"]
    for item in diff_items:
        lines.append(item.to_markdown())
        lines.append("")  # blank line between items

    lines.append("---\n🤔 **请点击下方按钮快捷操作，或者直接打字告诉我您的决定：**")

    # 核心修改：将按钮直接挂载在输出建议的这条消息上！
    msg = cl.Message(
        content="\n".join(lines),
        actions=[
            cl.Action(name="accept_all_action", payload={"action": "accept"}, label="✅ 全部接受"),
            cl.Action(name="reject_all_action", payload={"action": "reject"}, label="❌ 全部拒绝"),
        ]
    )
    await msg.send()
    cl.user_session.set("diff_msg", msg)


async def _execute_feedback(intent: str, rejected: List[int]) -> None:
    """执行校对反馈操作并输出文档"""
    cl.user_session.set(_KEY_STATE, "ready")  # 恢复状态

    # 移除界面上的按钮（无论是通过点击还是聊天触发该流程，都应令历史按钮消失）
    msg = cl.user_session.get("diff_msg")
    if msg:
        msg.actions = []
        await msg.update()
        cl.user_session.set("diff_msg", None)

    diff_items = cl.user_session.get(_KEY_DIFF_ITEMS, [])
    raw_issues = cl.user_session.get(_KEY_ISSUES, [])
    out_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES, b"")
    report = cl.user_session.get(_KEY_REPORT, {})
    filename = cl.user_session.get(_KEY_FILENAME, "output.docx")
    total = len(diff_items)

    if intent == "reject_all":
        await cl.Message(content="⏭️ 已跳过所有校对建议，正在生成最终文档…").send()
        await _provide_download(out_bytes, report, filename, applied=0)
        return

    if rejected:
        kept = total - len(rejected)
        await cl.Message(
            content=f"⏳ 已拒绝 **#{', #'.join(str(n) for n in sorted(rejected))}**，"
                    f"正在应用其余 **{kept}** 条建议…"
        ).send()
    else:
        await cl.Message(content=f"⏳ 正在应用全部 **{total}** 条校对建议…").send()

    try:
        from ui.diff_utils import apply_and_save_proofread
        final_bytes, applied = apply_and_save_proofread(
            out_bytes, raw_issues, excluded_numbers=rejected
        )
    except Exception as e:
        await cl.Message(content=f"⚠️ 应用校对建议出错: {e}").send()
        final_bytes, applied = out_bytes, 0

    await _provide_download(final_bytes, report, filename, applied=applied)


# ── General chat ────────────────────────────────────────────────────────────
'''  # 物理隔离版本的配套代码
async def _handle_chat(text: str) -> None:
    """处理用户输入：隔离 Slash 命令与普通聊天"""
    if not LLM_API_KEY:
        await cl.Message(
            content="💬 未配置 LLM API Key，暂无法进行对话或解析指令。",
            actions=_make_mode_actions(),
        ).send()
        return

    text_strip = text.strip()
    is_format_cmd = text_strip.startswith("/f ") or text_strip.startswith("/format ") or text_strip == "/f" or text_strip == "/format"

    # ════════════════════════════════════════════════════════════════════════
    # 分支 A：用户明确下达排版指令 (Slash Command)
    # ════════════════════════════════════════════════════════════════════════
    if is_format_cmd:
        # 提取真实的指令内容
        cmd_content = text_strip.replace("/format", "").replace("/f", "").strip()
        
        if not cmd_content:
            await cl.Message(content="⚠️ 请在命令后输入具体要求，例如：`/f 所有标题居中`").send()
            return
            
        thinking_msg = cl.Message(content="⏳ 正在将您的指令翻译为排版配置...")
        await thinking_msg.send()

        try:
            from agent.intent_parser import parse_formatting_request
            current_spec_path = cl.user_session.get(_KEY_SPEC_PATH, "specs/default.yaml")
            parsed = await parse_formatting_request(cmd_content, current_spec_path=current_spec_path)
            formatting_intent = parsed.get("overrides", {})
            routed_spec = parsed.get("spec_path", current_spec_path)
        except Exception as e:
            formatting_intent = None
            print(f"解析报错: {e}")

        if formatting_intent:
            current_overrides: Dict[str, Any] = cl.user_session.get(_KEY_SPEC_OVERRIDES, {})
            new_overrides = _deep_merge_dicts(copy.deepcopy(current_overrides), formatting_intent)
            cl.user_session.set(_KEY_SPEC_OVERRIDES, new_overrides)
            cl.user_session.set(_KEY_SPEC_PATH, routed_spec)

            pretty_intent = json.dumps(formatting_intent, ensure_ascii=False, indent=2)
            pretty_overrides = json.dumps(new_overrides, ensure_ascii=False, indent=2)

            input_bytes: bytes = cl.user_session.get(_KEY_INPUT_BYTES)
            if input_bytes:
                thinking_msg.content = (
                    f"✅ **指令已确认！**\n\n"
                    f"**增量修改：**\n```json\n{pretty_intent}\n```\n"
                    f"🚀 正在为您**重新生成文档**..."
                )
                await thinking_msg.update()

                filename: str = cl.user_session.get(_KEY_FILENAME, "document.docx")
                label_mode = cl.user_session.get(_KEY_LABEL_MODE, LLM_MODE)
                use_react = cl.user_session.get(_KEY_USE_REACT, False)
                max_iters = cl.user_session.get(_KEY_MAX_ITERS, REACT_MAX_ITERS)
                
                await _process_file(
                    input_bytes, filename, label_mode, use_react, max_iters,
                    overrides=new_overrides,
                    spec_path=routed_spec,
                )
            else:
                thinking_msg.content = (
                    f"✅ **已记录您的排版偏好！** 下次上传文档时将自动应用。\n\n"
                    f"**当前完整配置：**\n```json\n{pretty_overrides}\n```\n"
                    f"💡 请直接上传 `.docx` 文件。"
                )
                await thinking_msg.update()
        else:
            thinking_msg.content = "❌ 抱歉，未能从您的指令中提取出有效的排版属性，请换种说法重试。"
            await thinking_msg.update()
            
        return

    # ════════════════════════════════════════════════════════════════════════
    # 分支 B：普通自由交谈 (无需解析意图，速度极快)
    # ════════════════════════════════════════════════════════════════════════
    import openai as _openai
    history: List[dict] = cl.user_session.get(_KEY_CHAT_HISTORY, [])
    history.append({"role": "user", "content": text})

    try:
        client = _openai.AsyncOpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)
        msg = cl.Message(content="")
        await msg.send()
        reply_parts: List[str] = []
        
        async with await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是 MyAgent 文档格式化助手。"
                        "你可以帮助用户了解文档格式化知识、解答关于本工具的使用问题，也可以进行一般性的中文对话。"
                        "如果用户在聊天中提出了排版要求，请提醒他使用 '/f + 需求' 的命令格式。" # 顺便让大模型也知道这个规则
                    ),
                },
                *history[-_MAX_CHAT_HISTORY:],
            ],
            stream=True,
        ) as stream:
            async for chunk in stream:
                token = (chunk.choices[0].delta.content or "") if chunk.choices else ""
                if token:
                    reply_parts.append(token)
                    await msg.stream_token(token)
        await msg.update()
        
        reply = "".join(reply_parts)
        history.append({"role": "assistant", "content": reply})
        cl.user_session.set(_KEY_CHAT_HISTORY, history)
    except Exception as e:
        await cl.Message(content=f"💬 对话失败：{e}").send()
'''


def _is_format_command(text: str) -> bool:
    """判断输入是否为排版指令（以 /f 或 /format 开头）。"""
    t = text.strip()
    return t == "/f" or t == "/format" or t.startswith("/f ") or t.startswith("/format ")


def _extract_format_content(text: str) -> str:
    """从排版指令中提取实际内容（去掉前缀 /f 或 /format）。"""
    t = text.strip()
    if t.startswith("/format "):
        return t[len("/format "):].strip()
    if t.startswith("/f "):
        return t[len("/f "):].strip()
    return ""


def _is_review_command(text: str) -> bool:
    """判断输入是否为审阅指令（以 /r 或 /review 开头）。"""
    t = text.strip()
    return t == "/r" or t == "/review" or t.startswith("/r ") or t.startswith("/review ")


def _extract_review_content(text: str) -> str:
    """从审阅指令中提取实际内容（去掉前缀 /r 或 /review）。"""
    t = text.strip()
    if t.startswith("/review "):
        return t[len("/review "):].strip()
    if t.startswith("/r "):
        return t[len("/r "):].strip()
    return ""


# 仅用于"一般格式"检测的关键词（排除页眉/页脚/页码/目录类关键词）
_NON_HFT_FORMAT_KEYWORDS = [
    r"字体", r"宋体", r"黑体", r"楷体", r"仿宋",
    r"Times\s*New\s*Roman", r"Arial",
    r"字号", r"小四", r"三号", r"四号", r"五号", r"小三", r"小二",
    r"\d+\s*pt", r"\d+\s*磅",
    r"行距", r"行间距", r"倍行距",
    r"首行缩进", r"缩进",
    r"加粗", r"斜体", r"下划线",
    r"颜色", r"[黑红蓝绿白黄橙紫]色",
    r"对齐", r"两端对齐",
    r"段[前后]",
]


def _has_hft_intent(text: str) -> bool:
    """检查文本是否包含页眉/页脚/页码/目录相关关键词。"""
    from agent.intent_classifier import _match_any, _HEADER_FOOTER_TOC_KEYWORDS
    return _match_any(text, _HEADER_FOOTER_TOC_KEYWORDS)


def _has_non_hft_format_intent(text: str) -> bool:
    """检查文本是否包含非页眉/页脚/页码/目录的一般格式排版需求（颜色、字号、字体、行距等）。"""
    text_stripped = text.strip()
    for pattern in _NON_HFT_FORMAT_KEYWORDS:
        if re.search(pattern, text_stripped, re.IGNORECASE):
            return True
    return False


# ── New feature handlers ───────────────────────────────────────────────────

async def _handle_audit(doc_bytes: bytes) -> None:
    """Feature 3：对上传的文档进行排版一致性审阅，返回问题列表。"""
    import io
    from docx import Document as _Document
    from core.doc_audit import audit_document, format_audit_report
    from core.parser import parse_docx_to_blocks

    thinking_msg = cl.Message(content="🔍 正在分析文档排版一致性，请稍候...")
    await thinking_msg.send()

    try:
        doc_buf = io.BytesIO(doc_bytes)
        _, blocks = parse_docx_to_blocks(io.BytesIO(doc_bytes))
        doc = _Document(doc_buf)
        issues = audit_document(doc, blocks)
        report_md = format_audit_report(issues)
        thinking_msg.content = report_md
        await thinking_msg.update()
    except Exception as e:
        thinking_msg.content = f"❌ 审阅过程中出错：{e}"
        await thinking_msg.update()


async def _run_review_proofread(doc_bytes: bytes, filename: str) -> bool:
    """Run LLM proofreading on a document and show diff cards if issues are found.

    Returns True if diff cards were displayed (state set to 'awaiting_feedback'),
    or False if no issues were found / LLM is unavailable.
    """
    import io
    from core.parser import parse_docx_to_blocks
    from agent.llm_client import LLMClient, LLMCallError

    thinking_msg = cl.Message(content="🤖 正在进行 LLM 智能校对，检测错别字和标点问题...")
    await thinking_msg.send()

    try:
        _, blocks = parse_docx_to_blocks(io.BytesIO(doc_bytes))
        paragraphs = [b.text for b in blocks]

        client = LLMClient()
        proofread_result = await asyncio.to_thread(client.call_proofread, paragraphs)
        raw_issues = [issue.model_dump() for issue in proofread_result.issues]
        diff_items = build_diff_items(raw_issues)

        if diff_items:
            await thinking_msg.remove()
            cl.user_session.set(_KEY_OUTPUT_BYTES, doc_bytes)
            cl.user_session.set(_KEY_REPORT, {
                "meta": {
                    "paragraphs_before": len(paragraphs),
                    "paragraphs_after": len(paragraphs),
                }
            })
            cl.user_session.set(_KEY_FILENAME, filename)
            cl.user_session.set(_KEY_ISSUES, raw_issues)
            cl.user_session.set(_KEY_DIFF_ITEMS, diff_items)
            await _show_diff_cards(diff_items)
            cl.user_session.set(_KEY_STATE, "awaiting_feedback")
            return True
        else:
            thinking_msg.content = "✅ LLM 校对完成，未发现明显文本错误。"
            await thinking_msg.update()
            return False
    except LLMCallError as e:
        thinking_msg.content = f"⚠️ LLM 校对暂不可用：{e}"
        await thinking_msg.update()
        return False
    except Exception as e:
        thinking_msg.content = f"⚠️ LLM 校对出错：{e}"
        await thinking_msg.update()
        return False


async def _apply_review_flow(base_bytes: bytes, review_content: str, filename: str) -> None:
    """审阅流程公共入口：先审阅文档，若附带增量要求则只修改指定内容，最后进行 LLM 智能校对。

    :param base_bytes: 要审阅/修改的文档字节
    :param review_content: /r 命令后的增量要求（可为空字符串，表示纯审阅）
    :param filename: 文件名（用于下载链接）
    """
    # 步骤 1：始终执行文档排版一致性审阅
    await _handle_audit(base_bytes)

    # 步骤 2：若附带了增量排版要求，解析并只应用指定修改（不立即下载）
    _changes_applied = False
    if review_content:
        try:
            from agent.intent_parser import parse_review_request
            review_parsed = await parse_review_request(review_content)
            has_requirements = review_parsed.get("has_requirements", False)
            incremental_overrides = review_parsed.get("overrides", {})
            incremental_hft = review_parsed.get("hft_actions", {})
        except Exception as e:
            has_requirements = False
            incremental_overrides = {}
            incremental_hft = {}
            print(f"解析审阅增量意图异常: {e}")

        if has_requirements:
            if incremental_overrides:
                await _handle_partial_format(base_bytes, review_content, filename, _skip_download=True)
                _changes_applied = True
            if incremental_hft:
                current_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES) or base_bytes
                await _handle_header_footer_toc(
                    current_bytes, review_content, filename, pre_parsed_cmd=incremental_hft
                )
                _changes_applied = True

    # 步骤 3：LLM 智能校对（与 /f 流程一致，输出 diff 卡片供用户选择）
    if LLM_API_KEY:
        current_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES) or base_bytes
        proofread_shown = await _run_review_proofread(current_bytes, filename)
        if not proofread_shown and _changes_applied:
            # 无校对建议时，若有增量修改则提供下载
            out_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES)
            report = cl.user_session.get(_KEY_REPORT, {})
            await _provide_download(out_bytes, report, filename, applied=0)
    elif _changes_applied:
        # 未配置 LLM：若有增量修改则直接下载
        out_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES)
        report = cl.user_session.get(_KEY_REPORT, {})
        await _provide_download(out_bytes, report, filename, applied=0)


def _safe_float(value, default: float) -> float:
    """安全地将 value 转换为 float，忽略非数字字符（如"10.5pt"→10.5），失败时返回 default。"""
    if value is None:
        return default
    try:
        return float(value)
    except (TypeError, ValueError):
        # 尝试去掉非数字字符后再解析（如 "10.5pt"、"小五" 等）
        m = re.search(r"[\d.]+", str(value))
        if m:
            try:
                return float(m.group())
            except ValueError:
                pass
    return default


async def _handle_header_footer_toc(
    doc_bytes: bytes,
    user_text: str,
    filename: str,
    *,
    pre_parsed_cmd: dict | None = None,
) -> None:
    """处理页眉/页脚/页码/目录操作，返回修改后的文档。

    :param pre_parsed_cmd: 已由 LLM 解析好的命令字典（来自 parse_formatting_request 返回的
                           hft_actions）。若提供此参数，则跳过本地规则/LLM 二次解析步骤。
    """
    import io
    from docx import Document as _Document
    from core.header_footer_toc import (
        set_header, set_footer, add_page_numbers, format_toc_content,
        parse_header_footer_command, remove_header_border,
    )

    thinking_msg = cl.Message(content="⏳ 正在处理页眉/页脚/页码/目录...")
    await thinking_msg.send()

    try:
        doc = _Document(io.BytesIO(doc_bytes))

        if pre_parsed_cmd is not None:
            # 使用调用方已解析好的命令，跳过二次解析
            parsed_cmd = pre_parsed_cmd
        else:
            # 解析用户自然语言指令（本地规则优先）
            parsed_cmd = parse_header_footer_command(user_text)

            # 如果本地规则解析为空，尝试 LLM 解析
            if not parsed_cmd and LLM_API_KEY:
                try:
                    from agent.intent_parser import _extract_json
                    import openai as _openai
                    from config import LLM_BASE_URL, LLM_MODEL
                    client = _openai.AsyncOpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL, timeout=20.0)
                    resp = await client.chat.completions.create(
                        model=LLM_MODEL,
                        messages=[{
                            "role": "system",
                            "content": (
                                "你是文档页眉/页脚/页码/目录解析器。从用户指令中提取操作参数，"
                                "输出 JSON：{\"header\":{\"text\":\"...\"},\"footer\":{\"text\":\"...\"},"
                                "\"page_numbers\":{\"position\":\"footer\",\"start_at\":1,\"show_total\":false},"
                                "\"toc\":{\"title\":\"目录\"}}，仅包含用户提到的键。只输出 JSON。"
                            ),
                        }, {"role": "user", "content": user_text}],
                        temperature=0.1, max_tokens=300,
                    )
                    parsed_cmd = _extract_json(resp.choices[0].message.content.strip()) or {}
                except Exception as e:
                    print(f"LLM 页眉解析异常: {e}")

        if not parsed_cmd:
            thinking_msg.content = (
                "❓ 未能识别页眉/页脚/页码/目录指令。\n\n"
                "请更具体地描述，例如：\n"
                "- 「在页眉中写上'华中科技大学'，居中显示」\n"
                "- 「在页脚添加页码，从第1页开始」\n"
                "- 「在文档开头添加目录」"
            )
            await thinking_msg.update()
            return

        # 当前字体配置
        current_spec_path = cl.user_session.get(_KEY_SPEC_PATH, "specs/default.yaml")
        try:
            from core.spec import load_spec
            spec = load_spec(current_spec_path)
            zh_font = spec.raw.get("fonts", {}).get("zh", "宋体")
            en_font = spec.raw.get("fonts", {}).get("en", "Times New Roman")
        except Exception:
            zh_font, en_font = "宋体", "Times New Roman"

        actions_done = []

        # 页眉
        if "header" in parsed_cmd and parsed_cmd["header"]:
            h_cfg = parsed_cmd["header"]
            set_header(
                doc,
                text=h_cfg.get("text", ""),
                font_name_zh=zh_font,
                font_name_en=en_font,
                font_size_pt=_safe_float(h_cfg.get("font_size_pt"), 10.5),
                bold=bool(h_cfg.get("bold", False)),
                alignment=h_cfg.get("alignment", "center"),
            )
            actions_done.append(f"✅ 已设置页眉：「{h_cfg.get('text', '')}」")

        # 页脚
        if "footer" in parsed_cmd and parsed_cmd["footer"]:
            f_cfg = parsed_cmd["footer"]
            set_footer(
                doc,
                text=f_cfg.get("text", ""),
                font_name_zh=zh_font,
                font_name_en=en_font,
                font_size_pt=_safe_float(f_cfg.get("font_size_pt"), 10.5),
                alignment=f_cfg.get("alignment", "center"),
            )
            actions_done.append(f"✅ 已设置页脚：「{f_cfg.get('text', '')}」")

        # 页码
        if "page_numbers" in parsed_cmd and parsed_cmd["page_numbers"]:
            pn_cfg = parsed_cmd["page_numbers"]
            add_page_numbers(
                doc,
                position=pn_cfg.get("position", "footer"),
                alignment=pn_cfg.get("alignment", "center"),
                show_total=bool(pn_cfg.get("show_total", False)),
                font_name_zh=zh_font,
                font_name_en=en_font,
                start_at=pn_cfg.get("start_at"),
            )
            pos_label = "页眉" if pn_cfg.get("position") == "header" else "页脚"
            start_label = f"（从第 {pn_cfg['start_at']} 页开始）" if pn_cfg.get("start_at") else ""
            actions_done.append(f"✅ 已在{pos_label}中插入页码{start_label}")

        # 目录格式修改（修改已有目录的字体和字号，不插入新目录）
        if "toc_format" in parsed_cmd and isinstance(parsed_cmd["toc_format"], dict):
            fmt_cfg = parsed_cmd["toc_format"] or {}
            count = format_toc_content(
                doc,
                font_name_zh=fmt_cfg.get("font_name_zh", zh_font),
                font_name_en=fmt_cfg.get("font_name_en", en_font),
                font_size_pt=_safe_float(fmt_cfg.get("font_size_pt"), 12.0),
                bold_top_level=bool(fmt_cfg.get("bold_top_level", False)),
            )
            actions_done.append(f"✅ 已修改目录内容格式（共 {count} 个段落）")

        # 删除页眉横线
        if parsed_cmd.get("header_remove_border"):
            remove_header_border(doc)
            actions_done.append("✅ 已删除页眉横线")

        if not actions_done:
            thinking_msg.content = "❌ 抱歉，未能从您的指令中提取出有效的排版属性，请换种说法重试。"
            await thinking_msg.update()
            return

        # 保存修改后的文档到会话（供后续下载或校对使用）
        out_buf = io.BytesIO()
        doc.save(out_buf)
        out_bytes = out_buf.getvalue()
        cl.user_session.set(_KEY_OUTPUT_BYTES, out_bytes)

        thinking_msg.content = "\n".join(actions_done)
        await thinking_msg.update()

    except Exception as e:
        thinking_msg.content = f"❌ 处理页眉/页脚/页码/目录时出错：{e}"
        await thinking_msg.update()


async def _handle_partial_format(doc_bytes: bytes, user_text: str, filename: str, _skip_download: bool = False) -> None:
    """Feature 2：局部/定向排版——只应用用户指定的特定格式属性。

    :param _skip_download: 若为 True，则不输出下载链接（由调用方统一提供）。
    """
    import io
    from docx import Document as _Document
    from agent.intent_parser import parse_partial_format_request
    from core.partial_formatter import apply_partial_format

    thinking_msg = cl.Message(content="⏳ 正在解析局部排版指令...")
    await thinking_msg.send()

    try:
        # 解析要修改的属性
        parsed = await parse_partial_format_request(user_text)
        overrides = parsed.get("overrides", {})
        prop_name = parsed.get("property", "unknown")
        desc = parsed.get("description", "")

        if not overrides:
            thinking_msg.content = (
                "❓ 未能识别具体要修改的属性。\n\n"
                "请更具体地描述，例如：\n"
                "- 「只把正文行间距改为1.5倍」\n"
                "- 「只改正文字号为12pt」\n"
                "- 「只调整页面左边距为3cm」"
            )
            await thinking_msg.update()
            return

        doc = _Document(io.BytesIO(doc_bytes))
        report = apply_partial_format(doc, overrides)

        # 保存
        out_buf = io.BytesIO()
        doc.save(out_buf)
        out_bytes = out_buf.getvalue()
        cl.user_session.set(_KEY_OUTPUT_BYTES, out_bytes)

        counts = report.get("counts", {})
        counts_str = "、".join(f"{v} 个{k}段落" for k, v in counts.items() if k != "page_sections")
        page_str = f"，调整了 {counts.get('page_sections', 0)} 个页面节" if counts.get("page_sections") else ""

        thinking_msg.content = (
            f"✅ **局部排版完成！**\n\n"
            f"📝 修改项：{desc or prop_name}\n"
            f"📊 影响范围：{counts_str or '无段落变更'}{page_str}\n\n"
            "📥 文档已更新，其他格式保持不变。"
        )
        await thinking_msg.update()

        if not _skip_download:
            base_name = os.path.splitext(os.path.basename(filename))[0]
            out_el = cl.File(
                name=f"{base_name}_partial.docx",
                content=out_bytes,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            await cl.Message(content="📥 下载修改后的文档：", elements=[out_el]).send()

    except Exception as e:
        thinking_msg.content = f"❌ 局部排版处理出错：{e}"
        await thinking_msg.update()


async def _handle_locate_format(doc_bytes: bytes, user_text: str, filename: str) -> None:
    """Feature 4：定位文档中特定内容并重新排版，使其与周围格式一致。"""
    import io
    from docx import Document as _Document
    from agent.intent_parser import parse_locate_format_request
    from core.locate_formatter import locate_and_reformat

    thinking_msg = cl.Message(content="🔍 正在定位文档中的目标段落...")
    await thinking_msg.send()

    try:
        # 解析定位请求
        parsed = await parse_locate_format_request(user_text)
        locate_text = parsed.get("locate_text", "")
        format_action = parsed.get("format_action", "match_context")
        overrides = parsed.get("overrides", {})
        desc = parsed.get("description", "")

        if not locate_text:
            thinking_msg.content = (
                "❓ 未能从您的描述中提取定位关键词。\n\n"
                "请在消息中引用要定位的原文，例如：\n"
                "「'【大四上学期】全力冲刺目标...' 这部分和其他地方格式不同，帮我重新排版」"
            )
            await thinking_msg.update()
            return

        doc = _Document(io.BytesIO(doc_bytes))
        report = locate_and_reformat(doc, locate_text, format_action, overrides)

        if report["changed_count"] == 0 and not report["matched_paragraphs"]:
            thinking_msg.content = (
                f"🔍 {report.get('message', '')}\n\n"
                "💡 提示：请在消息中直接引用文档中的原文片段（无需完整引用，部分关键词即可）。"
            )
            await thinking_msg.update()
            return

        # 保存
        out_buf = io.BytesIO()
        doc.save(out_buf)
        out_bytes = out_buf.getvalue()
        cl.user_session.set(_KEY_OUTPUT_BYTES, out_bytes)

        matched_texts = [f"「{m['text'][:40]}...」" for m in report["matched_paragraphs"][:3]]
        matched_str = "\n".join(f"  - {t}" for t in matched_texts)

        applied_fmt = report.get("applied_format", {})
        fmt_desc_parts = []
        if "font_size_pt" in applied_fmt:
            fmt_desc_parts.append(f"字号 {applied_fmt['font_size_pt']}pt")
        if "line_spacing" in applied_fmt:
            fmt_desc_parts.append(f"行距 {applied_fmt['line_spacing']:.1f}倍")
        if "bold" in applied_fmt:
            fmt_desc_parts.append("加粗" if applied_fmt["bold"] else "取消加粗")
        if "alignment" in applied_fmt:
            fmt_desc_parts.append(f"对齐 {applied_fmt['alignment']}")
        fmt_str = "、".join(fmt_desc_parts) if fmt_desc_parts else "（与周围段落保持一致）"

        thinking_msg.content = (
            f"✅ **定位排版完成！**\n\n"
            f"📍 定位到 {len(report['matched_paragraphs'])} 个段落：\n{matched_str}\n\n"
            f"🎨 应用格式：{fmt_str}\n\n"
            f"📝 {report.get('message', '')}"
        )
        await thinking_msg.update()

        base_name = os.path.splitext(os.path.basename(filename))[0]
        out_el = cl.File(
            name=f"{base_name}_located.docx",
            content=out_bytes,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        await cl.Message(content="📥 下载修改后的文档：", elements=[out_el]).send()

    except Exception as e:
        thinking_msg.content = f"❌ 定位排版处理出错：{e}"
        await thinking_msg.update()


async def _handle_chat(text: str) -> None:
    """处理用户输入：通过指令前缀分发到各个处理分支。

    - 以 /f 或 /format 开头 → 排版指令（LLM 识别全部排版需求，包括页眉/页脚/目录等）
    - 以 /r 或 /review 开头 → 审阅指令（文档排版审阅 + 可选的增量排版）
    - 其他 → 普通聊天
    """
    if not LLM_API_KEY:
        await cl.Message(
            content="💬 未配置 LLM API Key，暂无法进行对话或解析指令。"
        ).send()
        return

    # ════════════════════════════════════════════════════════════════════════
    # 分支 A：排版指令（以 /f 或 /format 开头）
    # ════════════════════════════════════════════════════════════════════════
    if _is_format_command(text):
        cmd_content = _extract_format_content(text)

        if not cmd_content:
            await cl.Message(
                content=(
                    "⚠️ 请在命令后输入具体要求，例如：\n"
                    "> `/f 所有标题居中，正文字号12pt，在页脚添加页码`"
                )
            ).send()
            return

        thinking_msg = cl.Message(content="⏳ 正在将您的指令翻译为排版配置...")
        await thinking_msg.send()

        try:
            from agent.intent_parser import parse_formatting_request
            current_spec_path = cl.user_session.get(_KEY_SPEC_PATH, "specs/default.yaml")
            parsed = await parse_formatting_request(cmd_content, current_spec_path=current_spec_path)
            formatting_intent = parsed.get("overrides", {})
            hft_actions = parsed.get("hft_actions", {})
            routed_spec = parsed.get("spec_path", current_spec_path)
        except Exception as e:
            formatting_intent = None
            hft_actions = {}
            routed_spec = cl.user_session.get(_KEY_SPEC_PATH, "specs/default.yaml")
            print(f"解析排版意图异常: {e}")

        base_bytes: bytes = (
            cl.user_session.get(_KEY_OUTPUT_BYTES)
            or cl.user_session.get(_KEY_INPUT_BYTES)
        )
        filename: str = cl.user_session.get(_KEY_FILENAME, "document.docx")
        max_iters = cl.user_session.get(_KEY_MAX_ITERS, REACT_MAX_ITERS)

        has_spec_changes = bool(formatting_intent)
        has_hft_changes = bool(hft_actions)

        if not has_spec_changes and not has_hft_changes:
            thinking_msg.content = (
                "❌ 抱歉，未能从您的指令中提取出有效的排版属性，请换种说法重试。"
            )
            await thinking_msg.update()
            return

        if not base_bytes:
            # 无文档：记录偏好供下次上传使用（spec + HFT 部分同步暂存）
            msg_parts: List[str] = []

            if has_spec_changes:
                current_overrides: Dict[str, Any] = cl.user_session.get(_KEY_SPEC_OVERRIDES, {})
                new_overrides = _deep_merge_dicts(copy.deepcopy(current_overrides), formatting_intent)
                cl.user_session.set(_KEY_SPEC_OVERRIDES, new_overrides)
                cl.user_session.set(_KEY_SPEC_PATH, routed_spec)
                pretty_overrides = json.dumps(new_overrides, ensure_ascii=False, indent=2)
                msg_parts.append(
                    f"✅ **已记录您的排版偏好！** 下次上传文档时将自动应用。\n\n"
                    f"**当前完整配置：**\n```json\n{pretty_overrides}\n```"
                )

            if has_hft_changes:
                current_hft: Dict[str, Any] = cl.user_session.get(_KEY_HFT_ACTIONS, {})
                new_hft = {**current_hft, **hft_actions}
                cl.user_session.set(_KEY_HFT_ACTIONS, new_hft)
                pretty_hft = json.dumps(hft_actions, ensure_ascii=False, indent=2)
                msg_parts.append(
                    f"✅ **已记录页眉/页脚/页码/目录操作！** 下次上传文档时将自动应用。\n\n"
                    f"**HFT 配置：**\n```json\n{pretty_hft}\n```"
                )

            msg_parts.append("💡 请直接上传 `.docx` 文件。")
            thinking_msg.content = "\n\n".join(msg_parts)
            await thinking_msg.update()
            return

        # ── 有文档：先应用 spec 样式修改，再应用 HFT 操作 ──────────────────
        if has_spec_changes:
            current_overrides = cl.user_session.get(_KEY_SPEC_OVERRIDES, {})
            new_overrides = _deep_merge_dicts(copy.deepcopy(current_overrides), formatting_intent)
            cl.user_session.set(_KEY_SPEC_OVERRIDES, new_overrides)
            cl.user_session.set(_KEY_SPEC_PATH, routed_spec)

            pretty_intent = json.dumps(formatting_intent, ensure_ascii=False, indent=2)
            thinking_msg.content = (
                f"✅ **指令已确认！**\n\n"
                f"**增量修改：**\n```json\n{pretty_intent}\n```\n"
                f"📚 当前模板：`{routed_spec}`\n"
                f"🚀 正在对文档进行增量修改..."
            )
            await thinking_msg.update()

            await _process_file(
                base_bytes, filename, max_iters,
                overrides=new_overrides,
                spec_path=routed_spec,
                _hft_pending=has_hft_changes,
            )

        # ── 应用 HFT 操作（如页眉、页脚、页码、目录等）──────────────────────
        if has_hft_changes:
            # 使用最新处理结果（若 spec 已处理则用其输出，否则用原始输入）
            current_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES) or base_bytes
            await _handle_header_footer_toc(
                current_bytes, cmd_content, filename, pre_parsed_cmd=hft_actions
            )
            # 如果没有待处理的校对建议，提供下载
            if cl.user_session.get(_KEY_STATE) != "awaiting_feedback":
                out_bytes = cl.user_session.get(_KEY_OUTPUT_BYTES)
                report = cl.user_session.get(_KEY_REPORT, {})
                await _provide_download(out_bytes, report, filename, applied=0)

        return

    # ════════════════════════════════════════════════════════════════════════
    # 分支 B：审阅指令（以 /r 或 /review 开头）
    # ════════════════════════════════════════════════════════════════════════
    if _is_review_command(text):
        review_content = _extract_review_content(text)

        base_bytes = (
            cl.user_session.get(_KEY_OUTPUT_BYTES)
            or cl.user_session.get(_KEY_INPUT_BYTES)
        )
        if not base_bytes:
            await cl.Message(
                content=(
                    "📄 请先上传一个 `.docx` 文件，然后再使用审阅指令。\n\n"
                    "💡 上传文档后：\n"
                    "- `/r` — 审阅文档，列出排版问题\n"
                    "- `/r 把正文行距改为1.5倍` — 审阅 + 只做指定的增量修改"
                )
            ).send()
            return

        filename = cl.user_session.get(_KEY_FILENAME, "document.docx")
        await _apply_review_flow(base_bytes, review_content, filename)
        return

    # ════════════════════════════════════════════════════════════════════════
    # 分支 C：普通聊天（无 /f 或 /r 前缀，直接以 Structura 角色回复）
    # ════════════════════════════════════════════════════════════════════════
    import openai as _openai
    history: List[dict] = cl.user_session.get(_KEY_CHAT_HISTORY, [])
    history.append({"role": "user", "content": text})

    try:
        from config import LLM_BASE_URL, LLM_MODEL
        client = _openai.AsyncOpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL, timeout=30.0)
        msg = cl.Message(content="")
        await msg.send()
        reply_parts: List[str] = []

        async with await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是 Structura 文档格式化助手。"
                        "你可以帮助用户了解文档格式化知识、解答关于本工具的使用问题，也可以进行一般性的中文对话。"
                        "如果用户想对文档进行排版调整，请提醒他使用 `/f + 需求` 的命令格式。"
                        "如果用户想审阅文档排版问题或进行增量修改，请提醒他使用 `/r + 需求` 的命令格式。"
                    ),
                },
                *history[-_MAX_CHAT_HISTORY:],
            ],
            stream=True,
        ) as stream:
            async for chunk in stream:
                token = (chunk.choices[0].delta.content or "") if chunk.choices else ""
                if token:
                    reply_parts.append(token)
                    await msg.stream_token(token)
        await msg.update()

        reply = "".join(reply_parts)
        history.append({"role": "assistant", "content": reply})
        cl.user_session.set(_KEY_CHAT_HISTORY, history)
    except Exception as e:
        await cl.Message(content=f"💬 对话失败：{e}").send()
# ── User feedback handling ─────────────────────────────────────────────────

async def _handle_feedback(message: cl.Message) -> None:
    """Handle user's natural language response for the pending diff items."""
    text = message.content.strip()
    diff_items: List[DiffItem] = cl.user_session.get(_KEY_DIFF_ITEMS, [])
    total = len(diff_items)

    # 显示过渡动画
    thinking_msg = cl.Message(content="⏳ 正在理解您的处理决定...")
    await thinking_msg.send()

    # 🚀 召唤大模型解析意图！
    from agent.intent_parser import parse_feedback_intent
    result = await parse_feedback_intent(text, total)
    await thinking_msg.remove()

    intent = result.get("intent", "unknown")
    rejected = result.get("rejected_indices", [])

    if intent == "unknown":
        await cl.Message(
            content="❓ 没太听懂您的意思，请明确说明您想保留或拒绝哪些建议，或者直接点击上方的按钮哦。"
        ).send()
        return

    # 交给执行引擎
    await _execute_feedback(intent, rejected)

# ── Download helper ────────────────────────────────────────────────────────

async def _provide_download(
    out_bytes: bytes,
    report: dict,
    filename: str,
    *,
    applied: int,
) -> None:
    """Send download links for the output docx and report JSON."""
    meta = report.get("meta", {})
    para_before = meta.get("paragraphs_before", "?")
    para_after = meta.get("paragraphs_after", "?")

    summary_lines = [
        "✅ **处理完成！**",
        "",
        f"📊 段落数：{para_before} → {para_after}",
    ]
    if applied:
        summary_lines.append(f"✏️ 文本校对应用：{applied} 处")

    warnings_list = report.get("warnings", [])
    if warnings_list:
        summary_lines.append(f"⚠️ 警告：{len(warnings_list)} 条")

    await cl.Message(content="\n".join(summary_lines)).send()

    base_name = os.path.splitext(os.path.basename(filename))[0]
    output_el = cl.File(
        name=f"{base_name}_formatted.docx",
        content=out_bytes,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    report_el = cl.File(
        name=f"{base_name}_report.json",
        content=json.dumps(report, ensure_ascii=False, indent=2).encode("utf-8"),
        mime="application/json",
    )
    await cl.Message(
        content="📥 下载产物：",
        elements=[output_el, report_el],
    ).send()
