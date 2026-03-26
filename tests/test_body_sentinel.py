from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.docx_utils import iter_all_paragraphs
from core.formatter import apply_formatting
from core.parser import Block
from core.spec import load_spec


def _make_blocks_labels(doc: Document):
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(iter_all_paragraphs(doc))
    ]
    labels = {b.block_id: "body" for b in blocks}
    labels["_source"] = "test"
    return blocks, labels


def test_body_sentinel_scopes_formatting_to_body_only():
    spec = load_spec("specs/default.yaml")
    doc = Document()
    p_before = doc.add_paragraph("封面内容")
    p_before.paragraph_format.first_line_indent = Pt(10)
    doc.add_paragraph("{body}")
    p_body = doc.add_paragraph("这是正文段落。")
    doc.add_paragraph("{/body}")
    p_after = doc.add_paragraph("附录说明")
    p_after.paragraph_format.first_line_indent = Pt(11)

    blocks, labels = _make_blocks_labels(doc)
    report = apply_formatting(doc, blocks, labels, spec)

    expected_indent = float(spec.raw["body"]["first_line_chars"]) * float(spec.raw["body"]["font_size_pt"])
    assert abs((p_before.paragraph_format.first_line_indent.pt or 0) - 10.0) < 0.1
    assert abs((p_after.paragraph_format.first_line_indent.pt or 0) - 11.0) < 0.1
    assert abs((p_body.paragraph_format.first_line_indent.pt or 0) - expected_indent) < 0.1
    assert report["actions"]["body_sentinel_enabled"] is True
    assert report["actions"]["body_sentinel_scoped_paragraphs"] == 1


def test_inline_body_sentinel_removed_and_formatted():
    spec = load_spec("specs/default.yaml")
    doc = Document()
    p = doc.add_paragraph("{body}仅正文内容{/body}")

    blocks, labels = _make_blocks_labels(doc)
    report = apply_formatting(doc, blocks, labels, spec)

    expected_indent = float(spec.raw["body"]["first_line_chars"]) * float(spec.raw["body"]["font_size_pt"])
    assert "{body}" not in p.text
    assert "{/body}" not in p.text
    assert p.text == "仅正文内容"
    assert abs((p.paragraph_format.first_line_indent.pt or 0) - expected_indent) < 0.1
    assert report["actions"]["body_sentinel_enabled"] is True
