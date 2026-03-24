from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.parser import Block
from core.spec import load_spec
from core.formatter import apply_formatting


def _mk_blocks(doc):
    blocks = []
    for i, p in enumerate(doc.paragraphs):
        blocks.append(Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i))
    return blocks


def test_cover_paragraph_skipped_and_section_roles_applied():
    doc = Document()
    doc.add_paragraph("课程报告")
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论........1")
    doc.add_paragraph("课程要求")
    doc.add_paragraph("按模板完成正文，注意行距。")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这是正文内容。")

    blocks = _mk_blocks(doc)
    labels = {b.block_id: "body" for b in blocks}
    labels["_source"] = "test"
    spec = load_spec("specs/default.yaml")

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["formatted"]["counts"].get("cover_skipped", 0) >= 1
    assert report["formatted"]["counts"].get("toc", 0) >= 1
    assert report["formatted"]["counts"].get("requirement", 0) >= 1


def test_caption_and_table_centering_actions_reported():
    doc = Document()
    p_img = doc.add_paragraph("图片占位段")

    p_caption = doc.add_paragraph("图1 示例图片")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "单元格"

    blocks = _mk_blocks(doc)
    labels = {b.block_id: "body" for b in blocks}
    labels[2] = "caption"
    labels["_source"] = "test"
    spec = load_spec("specs/default.yaml")

    import core.formatter as fmt
    orig = fmt._paragraph_has_inline_drawing
    fmt._paragraph_has_inline_drawing = lambda p: (p.text or "").strip() == "图片占位段"
    try:
        report = apply_formatting(doc, blocks, labels, spec)
    finally:
        fmt._paragraph_has_inline_drawing = orig

    assert p_img.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p_caption.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert report["actions"]["media_paragraphs_centered"] >= 1
    assert report["actions"]["tables_centered"] >= 1
