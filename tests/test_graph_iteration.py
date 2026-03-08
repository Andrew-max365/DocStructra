# tests/test_graph_iteration.py
"""Tests for graph iteration stopping conditions using mocked LangGraph nodes."""
from __future__ import annotations

from unittest.mock import MagicMock, patch

import pytest

from agent.graph.nodes import retry_router
from agent.graph.react_schemas import GraphState


def _make_state(**overrides) -> GraphState:
    base: GraphState = {
        "input_path": "in.docx",
        "output_path": "out.docx",
        "spec_path": "specs/default.yaml",
        "label_mode": "unified",
        "max_iters": 3,
        "current_iter": 0,
        "thoughts": [],
        "actions": [],
        "observations": [],
        "errors": [],
        "passed": False,
        "finished": False,
        "report": {},
        "blocks": None,
        "labels": None,
        "doc": None,
        "overrides": None,
        "needs_llm": False,
        "triggered_indices": [],
        "hybrid_triggers": {},
        "proofread_issues": [],
    }
    base.update(overrides)
    return base


# --- retry_router tests ---

def test_retry_router_end_when_passed():
    state = _make_state(passed=True, current_iter=1)
    assert retry_router(state) == "end"


def test_retry_router_end_when_max_iters_reached():
    state = _make_state(passed=False, current_iter=3, max_iters=3)
    assert retry_router(state) == "end"


def test_retry_router_reason_when_not_passed():
    state = _make_state(passed=False, current_iter=1, max_iters=3)
    assert retry_router(state) == "reason"


def test_retry_router_end_when_iter_exceeds_max():
    state = _make_state(passed=False, current_iter=5, max_iters=3)
    assert retry_router(state) == "end"


# --- reason_node tests ---

def test_reason_node_increments_iter():
    from agent.graph.nodes import reason_node
    state = _make_state(current_iter=0, errors=[])
    result = reason_node(state)
    assert result["current_iter"] == 1
    assert len(result["thoughts"]) == 1
    assert len(result["actions"]) == 1


def test_reason_node_appends_to_existing():
    from agent.graph.nodes import reason_node
    state = _make_state(
        current_iter=1,
        thoughts=["first thought"],
        actions=[[{"action_type": "no_op", "block_id": -1, "params": {}}]],
        errors=["some error"],
    )
    result = reason_node(state)
    assert result["current_iter"] == 2
    assert len(result["thoughts"]) == 2
    # reason_node no longer echoes prior state errors; it reports its own LLM call result
    assert isinstance(result["thoughts"][1], str)


def test_reason_node_produces_no_op_action():
    from agent.graph.nodes import reason_node
    state = _make_state(current_iter=0)
    result = reason_node(state)
    actions = result["actions"][-1]
    assert len(actions) == 1
    assert actions[0]["action_type"] == "no_op"


# --- act_node tests ---

def test_act_node_applies_set_role(tmp_path):
    """act_node should override label when action_type=set_role."""
    from agent.graph.nodes import act_node
    from docx import Document
    from core.parser import parse_docx_to_blocks
    from core.judge import rule_based_labels

    doc = Document()
    doc.add_paragraph("Title paragraph")
    doc.add_paragraph("Body paragraph")
    docx_path = str(tmp_path / "test.docx")
    doc.save(docx_path)
    doc2, blocks = parse_docx_to_blocks(docx_path)
    initial_labels = rule_based_labels(blocks, doc=doc2)

    state = _make_state(
        doc=doc2,
        blocks=blocks,
        labels=initial_labels,
        actions=[[{"action_type": "set_role", "block_id": 0, "params": {"role": "h1"}}]],
    )
    result = act_node(state)
    assert result["labels"][0] == "h1"


def test_act_node_applies_fix_heading_level(tmp_path):
    from agent.graph.nodes import act_node
    from docx import Document
    from core.parser import parse_docx_to_blocks
    from core.judge import rule_based_labels

    doc = Document()
    doc.add_paragraph("Section heading")
    docx_path = str(tmp_path / "test2.docx")
    doc.save(docx_path)
    doc2, blocks = parse_docx_to_blocks(docx_path)
    initial_labels = rule_based_labels(blocks, doc=doc2)

    state = _make_state(
        doc=doc2,
        blocks=blocks,
        labels=initial_labels,
        actions=[[{"action_type": "fix_heading_level", "block_id": 0, "params": {"level": 2}}]],
    )
    result = act_node(state)
    assert result["labels"][0] == "h2"


# --- trigger_node adaptive triggers tests ---
def test_trigger_node_empty_paragraphs(tmp_path):
    """测试当空段落占比过高时触发 LLM 清理建议"""
    from agent.graph.nodes import trigger_node
    from docx import Document
    from core.parser import parse_docx_to_blocks
    from core.judge import rule_based_labels
    
    doc = Document()
    doc.add_paragraph("正常段落 1")
    doc.add_paragraph("")  # 空段落
    doc.add_paragraph("正常段落 2")
    doc.add_paragraph("   ") # 纯空格也是空段落
    doc.add_paragraph("")  # 再次空段落
    # 5 个段落，3 个空的 -> 60% > 30% 阈值
    
    docx_path = str(tmp_path / "test_empty.docx")
    doc.save(docx_path)
    doc2, blocks = parse_docx_to_blocks(docx_path)
    initial_labels = rule_based_labels(blocks, doc=doc2)
    
    state = _make_state(
        doc=doc2, blocks=blocks, labels=initial_labels
    )
    result = trigger_node(state)
    
    assert result["needs_llm"] is True
    reasons_str = " ".join(result["hybrid_triggers"]["reasons"])
    assert "空段落占比过高" in reasons_str

def test_trigger_node_font_size_variance(tmp_path):
    """测试同一种类的角色出现多种字号时触发警告"""
    from agent.graph.nodes import trigger_node
    from docx import Document
    from docx.shared import Pt
    from core.parser import parse_docx_to_blocks
    from core.judge import rule_based_labels
    
    doc = Document()
    # 创建4个段落，人为指定不同的 font_size，让它们都被识别为 body
    sizes = [10, 11, 12, 14]
    for size in sizes:
        p = doc.add_paragraph()
        run = p.add_run(f"字号测试段落 {size}")
        run.font.size = Pt(size)
        
    docx_path = str(tmp_path / "test_fonts.docx")
    doc.save(docx_path)
    doc2, blocks = parse_docx_to_blocks(docx_path)
    initial_labels = rule_based_labels(blocks, doc=doc2)
    
    state = _make_state(
        doc=doc2, blocks=blocks, labels=initial_labels
    )
    result = trigger_node(state)
    
    # 字号校验目前只是挂 alert append reasons，如果段落少没其他特征，整体可能还是 doesn't need LLM
    # 但 reasons 必须包含字号不统一的警告
    reasons_str = " ".join(result["hybrid_triggers"]["reasons"])
    assert "存在多种不同字号" in reasons_str
