from io import BytesIO

from docx import Document

from core.awdp import (
    AWDPValidationError,
    get_awdp_prompt_template,
    render_awdp_markdown_to_docx_bytes,
    validate_awdp_markdown,
)


def _valid_awdp_markdown() -> str:
    return """---
protocol: AWDP-1.0
title: 示例文档
lang: zh-CN
---
# 一级标题

这是第一段正文。

这是第二段正文。

## 二级标题

| 列1 | 列2 |
| --- | --- |
| A | B |

```python
print("hello")
```

![流程图](https://example.com/img.png "系统流程图")
"""


def test_prompt_template_contains_protocol_and_rules():
    prompt = get_awdp_prompt_template()
    assert "AWDP-1.0" in prompt
    assert "YAML Front Matter" in prompt
    assert "只允许三层标题" in prompt


def test_validate_awdp_markdown_ok():
    parsed = validate_awdp_markdown(_valid_awdp_markdown())
    assert parsed.front_matter["protocol"] == "AWDP-1.0"
    assert parsed.body.startswith("# 一级标题")


def test_validate_awdp_markdown_rejects_missing_caption_and_code_lang():
    bad = """---
protocol: AWDP-1.0
---
# 标题

``` 
no lang
```

![img](https://example.com/x.png)
"""
    try:
        validate_awdp_markdown(bad)
        assert False, "expected validation error"
    except AWDPValidationError as e:
        assert any("代码块缺少语言声明" in msg for msg in e.errors)
        assert any("图片缺少标题" in msg for msg in e.errors)


def test_render_awdp_markdown_to_docx_bytes_generates_document():
    out = render_awdp_markdown_to_docx_bytes(_valid_awdp_markdown())
    assert isinstance(out, bytes)
    assert len(out) > 0

    doc = Document(BytesIO(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("一级标题" in t for t in texts)
    assert any("第一段正文" in t for t in texts)


def test_multiline_paragraph_passes_validation():
    """Consecutive lines within the same paragraph must not cause validation errors."""
    md = """---
protocol: AWDP-1.0
title: 多行段落测试
---
# 标题

这是第一行，
这是第二行，
这是第三行，属于同一段落。

下一段。
"""
    parsed = validate_awdp_markdown(md)
    assert parsed.front_matter["protocol"] == "AWDP-1.0"


def test_multiline_paragraph_merged_in_render():
    """Consecutive lines belonging to the same paragraph must be merged into one."""
    md = """---
protocol: AWDP-1.0
title: 合并测试
---
# 标题

第一行内容，
第二行内容。
"""
    out = render_awdp_markdown_to_docx_bytes(md)
    doc = Document(BytesIO(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Both lines should appear in a single merged paragraph
    assert any("第一行内容" in t and "第二行内容" in t for t in texts), \
        f"Expected merged paragraph, got: {texts}"


def test_inline_formatting_bold_italic_code():
    """**bold**, *italic*, and `code` must produce properly formatted runs."""
    md = """---
protocol: AWDP-1.0
title: 内联格式测试
---
# 标题

这里有**加粗**和*斜体*还有`代码`。
"""
    out = render_awdp_markdown_to_docx_bytes(md)
    doc = Document(BytesIO(out))

    # Find the paragraph with inline formatting
    formatted_para = next(
        (p for p in doc.paragraphs if "加粗" in p.text),
        None,
    )
    assert formatted_para is not None, "Could not find paragraph with inline formatting"

    bold_runs = [r for r in formatted_para.runs if "加粗" in r.text and r.bold]
    italic_runs = [r for r in formatted_para.runs if "斜体" in r.text and r.italic]
    code_runs = [r for r in formatted_para.runs if "代码" in r.text and r.font.name == "Courier New"]
    assert bold_runs, "Expected bold run containing '加粗'"
    assert italic_runs, "Expected italic run containing '斜体'"
    assert code_runs, "Expected code run (Courier New) containing '代码'"


def test_table_has_grid_style():
    """Rendered tables must use Table Grid style."""
    md = """---
protocol: AWDP-1.0
title: 表格测试
---
# 标题

| 列1 | 列2 |
| --- | --- |
| A | B |
"""
    out = render_awdp_markdown_to_docx_bytes(md)
    doc = Document(BytesIO(out))
    assert doc.tables, "Expected at least one table"
    assert doc.tables[0].style.name == "Table Grid"


def test_crlf_line_endings_accepted():
    """Windows CRLF line endings must be normalised and accepted."""
    md = "---\r\nprotocol: AWDP-1.0\r\ntitle: CRLF测试\r\n---\r\n# 标题\r\n\r\n正文。\r\n"
    out = render_awdp_markdown_to_docx_bytes(md)
    assert isinstance(out, bytes) and len(out) > 0

    doc = Document(BytesIO(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("标题" in t for t in texts)
    assert any("正文" in t for t in texts)

