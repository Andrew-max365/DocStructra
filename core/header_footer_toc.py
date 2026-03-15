# core/header_footer_toc.py
"""
页眉、页脚、页码、目录相关排版工具函数。

提供的功能：
  - set_header(doc, text, ...)         : 设置页眉文本及格式
  - set_footer(doc, text, ...)         : 设置页脚文本及格式
  - add_page_numbers(doc, ...)         : 在页眉或页脚中插入页码域
  - insert_toc(doc, ...)              : 在文档头部插入目录
"""
from __future__ import annotations

import re
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.docx_utils import set_run_fonts


# ─────────────────────────────────────────────────────────────────────────────
# 内部工具
# ─────────────────────────────────────────────────────────────────────────────

def _resolve_alignment(align_str: str) -> WD_ALIGN_PARAGRAPH:
    mapping = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get((align_str or "center").strip().lower(), WD_ALIGN_PARAGRAPH.CENTER)


def _clear_paragraph(paragraph) -> None:
    """清空段落中所有 run（保留段落本身）。"""
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)


def _add_fld_char(run_elem, fld_type: str) -> None:
    """在 run XML 元素中添加 w:fldChar 节点。"""
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), fld_type)
    run_elem.append(fld)


def _add_instr_text(run_elem, instr: str) -> None:
    """在 run XML 元素中添加 w:instrText 节点。"""
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run_elem.append(instr_el)


def _insert_page_num_field(paragraph, instr: str = "PAGE", rPr_el=None) -> None:
    """
    在 paragraph 中插入 Word 页码域（PAGE / NUMPAGES 等）。
    生成格式为：{ PAGE } 或 { PAGE } / { NUMPAGES } 等。
    可传入 rPr_el 以应用字体/字号/粗体等格式到域的每个 run。
    """
    # begin
    r1 = OxmlElement("w:r")
    if rPr_el is not None:
        r1.append(deepcopy(rPr_el))
    _add_fld_char(r1, "begin")
    paragraph._p.append(r1)
    # instrText
    r2 = OxmlElement("w:r")
    if rPr_el is not None:
        r2.append(deepcopy(rPr_el))
    _add_instr_text(r2, f" {instr} ")
    paragraph._p.append(r2)
    # end
    r3 = OxmlElement("w:r")
    if rPr_el is not None:
        r3.append(deepcopy(rPr_el))
    _add_fld_char(r3, "end")
    paragraph._p.append(r3)


def _build_rPr_element(
    font_name_zh: str,
    font_name_en: str,
    font_size_pt: float,
    bold: bool,
    italic: bool,
) -> "OxmlElement":
    """构建 w:rPr XML 元素，包含字体/字号/粗体/斜体设置。"""
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name_en)
    rFonts.set(qn("w:hAnsi"), font_name_en)
    rFonts.set(qn("w:eastAsia"), font_name_zh)
    rFonts.set(qn("w:cs"), font_name_en)
    rPr.append(rFonts)
    sz_val = str(int(font_size_pt * 2))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), sz_val)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), sz_val)
    rPr.append(sz)
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    return rPr


# 页码格式 fmt 到 Word 域开关的映射
_FMT_SWITCHES = {
    "arabic": "",
    "roman": r" \* roman",
    "ROMAN": r" \* ROMAN",
    "alpha": r" \* alphabetic",
    "ALPHA": r" \* ALPHABETIC",
}


def _apply_run_format(
    run,
    font_name_zh: str = "宋体",
    font_name_en: str = "Times New Roman",
    font_size_pt: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color_hex: Optional[str] = None,
) -> None:
    """统一设置 run 字体/字号/粗斜体/颜色。"""
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        try:
            color_hex = color_hex.lstrip("#")
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16)
            b = int(color_hex[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        except Exception:
            pass
    set_run_fonts(run, zh_font=font_name_zh, en_font=font_name_en)


# ─────────────────────────────────────────────────────────────────────────────
# 公共接口
# ─────────────────────────────────────────────────────────────────────────────

def set_header(
    doc: Document,
    text: str,
    *,
    font_name_zh: str = "宋体",
    font_name_en: str = "Times New Roman",
    font_size_pt: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    alignment: str = "center",
    color_hex: Optional[str] = None,
    section_index: int = 0,
) -> None:
    """
    设置文档指定节的页眉文本及格式。

    :param doc:           python-docx Document 对象
    :param text:          页眉文本（如"华中科技大学"）
    :param font_name_zh:  中文字体
    :param font_name_en:  英文字体
    :param font_size_pt:  字号（磅）
    :param bold:          是否加粗
    :param italic:        是否斜体
    :param alignment:     对齐方式 center/left/right
    :param color_hex:     字体颜色（十六进制，如 "FF0000"）
    :param section_index: 节序号（默认 0，即第一节）
    """
    sections = doc.sections
    if section_index >= len(sections):
        section_index = 0
    sec = sections[section_index]
    sec.different_first_page_header_footer = False
    header = sec.header
    header.is_linked_to_previous = False

    # 取第一个段落（若无则新建）
    if header.paragraphs:
        para = header.paragraphs[0]
        _clear_paragraph(para)
    else:
        para = header.add_paragraph()

    para.alignment = _resolve_alignment(alignment)
    run = para.add_run(text)
    _apply_run_format(run, font_name_zh, font_name_en, font_size_pt, bold, italic, color_hex)


def set_footer(
    doc: Document,
    text: str,
    *,
    font_name_zh: str = "宋体",
    font_name_en: str = "Times New Roman",
    font_size_pt: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    alignment: str = "center",
    color_hex: Optional[str] = None,
    section_index: int = 0,
) -> None:
    """
    设置文档指定节的页脚文本及格式。
    参数含义同 set_header。
    """
    sections = doc.sections
    if section_index >= len(sections):
        section_index = 0
    sec = sections[section_index]
    footer = sec.footer
    footer.is_linked_to_previous = False

    if footer.paragraphs:
        para = footer.paragraphs[0]
        _clear_paragraph(para)
    else:
        para = footer.add_paragraph()

    para.alignment = _resolve_alignment(alignment)
    run = para.add_run(text)
    _apply_run_format(run, font_name_zh, font_name_en, font_size_pt, bold, italic, color_hex)


def add_page_numbers(
    doc: Document,
    *,
    position: str = "footer",
    alignment: str = "center",
    fmt: str = "arabic",
    show_total: bool = False,
    separator: str = "/",
    font_name_zh: str = "宋体",
    font_name_en: str = "Times New Roman",
    font_size_pt: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    section_index: int = 0,
    start_at: Optional[int] = None,
) -> None:
    """
    在页眉或页脚中插入页码域。

    :param doc:           Document 对象
    :param position:      "header" 或 "footer"
    :param alignment:     对齐 center/left/right
    :param fmt:           页码样式 "arabic"（阿拉伯数字）/ "roman" / "alpha"
    :param show_total:    是否同时显示总页数（如"第 3 页 / 共 10 页"）
    :param separator:     页码与总页数的分隔符
    :param font_name_zh:  中文字体
    :param font_name_en:  英文字体
    :param font_size_pt:  字号
    :param bold:          加粗
    :param italic:        斜体
    :param section_index: 节序号
    :param start_at:      从哪一页开始编号（设置节起始页码）
    """
    sections = doc.sections
    if section_index >= len(sections):
        section_index = 0
    sec = sections[section_index]

    # 设置起始页码（通过 XML 属性直接设置）
    if start_at is not None:
        pgNumType = sec._sectPr.find(qn("w:pgNumType"))
        if pgNumType is None:
            pgNumType = OxmlElement("w:pgNumType")
            sec._sectPr.append(pgNumType)
        pgNumType.set(qn("w:start"), str(start_at))

    # 获取页眉或页脚
    if position.lower() == "header":
        container = sec.header
        container.is_linked_to_previous = False
    else:
        container = sec.footer
        container.is_linked_to_previous = False

    if container.paragraphs:
        para = container.paragraphs[0]
        _clear_paragraph(para)
    else:
        para = container.add_paragraph()

    para.alignment = _resolve_alignment(alignment)

    rpr_kwargs = dict(
        font_name_zh=font_name_zh,
        font_name_en=font_name_en,
        font_size_pt=font_size_pt,
        bold=bold,
        italic=italic,
    )

    # 构建域 run 的 rPr，用于给 PAGE / NUMPAGES 字段应用相同的字体格式
    rPr_el = _build_rPr_element(font_name_zh, font_name_en, font_size_pt, bold, italic)

    # 页码格式开关（arabic 不需要额外开关）
    fmt_switch = _FMT_SWITCHES.get(fmt, "")
    page_instr = f"PAGE{fmt_switch}"
    numpage_instr = f"NUMPAGES{fmt_switch}"

    if show_total:
        # 格式：第 {PAGE} 页 / 共 {NUMPAGES} 页
        r_prefix = para.add_run("第 ")
        _apply_run_format(r_prefix, **rpr_kwargs)
        _insert_page_num_field(para, page_instr, rPr_el)
        r_mid = para.add_run(f" 页 {separator} 共 ")
        _apply_run_format(r_mid, **rpr_kwargs)
        _insert_page_num_field(para, numpage_instr, rPr_el)
        r_suffix = para.add_run(" 页")
        _apply_run_format(r_suffix, **rpr_kwargs)
    else:
        # 简单页码：{PAGE}，应用字体格式
        _insert_page_num_field(para, page_instr, rPr_el)


def _configure_toc_styles(
    doc: Document,
    font_name_zh: str,
    font_name_en: str,
    font_size_pt: float,
    bold_top_level: bool,
) -> None:
    """
    修改或创建文档中的 TOC1/TOC2/TOC3 段落样式，应用目录内容字体格式。
    Word 会在更新目录时将样式应用到各级条目。
    """
    for level in range(1, 4):
        # Word 内置样式名为 "toc N"（小写+空格，python-docx 别名兼容）
        style = None
        for try_name in (f"toc {level}", f"TOC {level}", f"TOC{level}"):
            try:
                style = doc.styles[try_name]
                break
            except KeyError:
                continue
        if style is None:
            try:
                style = doc.styles.add_style(f"toc {level}", WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                continue
        try:
            style.font.size = Pt(font_size_pt)
            style.font.bold = (bold_top_level and level == 1)
            # 修改样式级别的 rFonts（中英文字体分离）
            rPr = style.element.get_or_add_rPr()
            existing_rFonts = rPr.find(qn("w:rFonts"))
            if existing_rFonts is not None:
                rPr.remove(existing_rFonts)
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), font_name_en)
            rFonts.set(qn("w:hAnsi"), font_name_en)
            rFonts.set(qn("w:eastAsia"), font_name_zh)
            rFonts.set(qn("w:cs"), font_name_en)
            rPr.append(rFonts)
        except Exception:
            pass


def insert_toc(
    doc: Document,
    *,
    title: str = "目录",
    title_font_name_zh: str = "黑体",
    title_font_name_en: str = "Times New Roman",
    title_font_size_pt: float = 18.0,
    title_bold: bool = True,
    title_alignment: str = "center",
    content_font_name_zh: str = "宋体",
    content_font_name_en: str = "Times New Roman",
    content_font_size_pt: float = 12.0,
    content_bold_top_level: bool = True,
    insert_position: int = 0,
) -> None:
    """
    在文档指定位置插入目录（TOC）。
    插入的目录依赖 Word 的自动目录域（TOC 域代码），需要在 Word 中手动更新或
    通过 doc.update_fields() 触发更新（某些环境支持）。

    :param doc:                   Document 对象
    :param title:                 目录标题（默认"目录"）
    :param title_font_name_zh:    目录标题中文字体（默认黑体）
    :param title_font_name_en:    目录标题英文字体
    :param title_font_size_pt:    目录标题字号（默认 18pt = 小二）
    :param title_bold:            目录标题是否加粗
    :param title_alignment:       目录标题对齐方式
    :param content_font_name_zh:  目录内容中文字体（默认宋体）
    :param content_font_name_en:  目录内容英文字体
    :param content_font_size_pt:  目录内容字号（默认 12pt = 小四）
    :param content_bold_top_level: 一级目录条目是否加粗
    :param insert_position:       插入位置（段落序号，0 = 最前面）
    """
    # ── 1. 目录标题段落 ──────────────────────────────────────────────────────
    title_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), title_alignment.lower() if title_alignment else "center")
    pPr.append(jc)
    title_para.append(pPr)

    title_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # 字体
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:hint"), "eastAsia")
    rFonts.set(qn("w:ascii"), title_font_name_en)
    rFonts.set(qn("w:eastAsia"), title_font_name_zh)
    rFonts.set(qn("w:hAnsi"), title_font_name_en)
    rPr.append(rFonts)
    # 字号（单位：半磅 = pt * 2）
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(title_font_size_pt * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(title_font_size_pt * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    if title_bold:
        rPr.append(OxmlElement("w:b"))
    title_run.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = title
    title_run.append(t_el)
    title_para.append(title_run)

    # ── 2. TOC 域段落 ────────────────────────────────────────────────────────
    toc_para = OxmlElement("w:p")
    # 目录内容的字体格式通过 rPr 控制
    toc_pPr = OxmlElement("w:pPr")
    toc_pStyle = OxmlElement("w:pStyle")
    toc_pStyle.set(qn("w:val"), "TOC1")
    toc_pPr.append(toc_pStyle)
    toc_para.append(toc_pPr)

    # begin
    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    toc_para.append(r_begin)

    # instrText：TOC \o "1-3" \h \z \u
    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr.append(instr)
    toc_para.append(r_instr)

    # end
    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    toc_para.append(r_end)

    # ── 3. 插入到文档 ────────────────────────────────────────────────────────
    body = doc.element.body
    all_paras = body.findall(qn("w:p"))

    if insert_position <= 0 or not all_paras:
        # 插在最前面
        if all_paras:
            ref = all_paras[0]
            body.insert(list(body).index(ref), toc_para)
            body.insert(list(body).index(ref), title_para)
        else:
            body.append(title_para)
            body.append(toc_para)
    else:
        idx = min(insert_position, len(all_paras)) - 1
        ref = all_paras[idx] if idx < len(all_paras) else all_paras[-1]
        ref_idx = list(body).index(ref) + 1
        body.insert(ref_idx, toc_para)
        body.insert(ref_idx, title_para)

    # ── 4. 配置目录内容样式（TOC1/TOC2/TOC3）────────────────────────────────
    _configure_toc_styles(
        doc,
        font_name_zh=content_font_name_zh,
        font_name_en=content_font_name_en,
        font_size_pt=content_font_size_pt,
        bold_top_level=content_bold_top_level,
    )


# ─────────────────────────────────────────────────────────────────────────────
# 自然语言解析辅助（供 UI 层调用，无需 LLM）
# ─────────────────────────────────────────────────────────────────────────────

def parse_header_footer_command(text: str) -> dict:
    """
    从用户自然语言文本中解析页眉/页脚/页码/目录命令参数。
    返回 dict，可能包含以下键：
      "header": {"text": ..., "alignment": ..., "font_size_pt": ...}
      "footer": {"text": ..., "alignment": ...}
      "page_numbers": {"position": "footer"|"header", "start_at": int, "show_total": bool}
      "toc": {"title": "目录", "insert_position": int}
    """
    result = {}

    # ── 页眉 ─────────────────────────────────────────────────────────────────
    _Q = r"""[「『\u201c\u2018'"]"""   # opening quotes (incl. ASCII ' and ")
    _QC = r"""[」』\u201d\u2019'"]"""  # closing quotes
    header_patterns = [
        r'页眉.*?' + _Q + r'(.*?)' + _QC,
        r'页眉.*?(?:内容|文字|写|加|设置|为)[：:]\s*(.+)',
        r'页眉.*?(?:显示|加上|写上|改为)\s*' + _Q + r'(.*?)' + _QC,
        r'(?:设置|添加|增加)页眉.*?' + _Q + r'(.*?)' + _QC,
        r'页眉[上中]?(?:写|加|为|显示)\s*' + _Q + r'?([\u4e00-\u9fffA-Za-z0-9\s]+)' + _QC + r'?',
    ]
    for pat in header_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            result["header"] = {"text": m.group(1).strip()}
            break

    # 如果没有匹配到引号内容，尝试宽松匹配
    if "header" not in result and re.search(r"页眉", text):
        m = re.search(r"页眉.*?[：:]\s*(.+?)(?:[，,。！\n]|$)", text)
        if m:
            result["header"] = {"text": m.group(1).strip()}

    # 检测字体要求
    if "header" in result:
        align_m = re.search(r"(居中|左对齐|右对齐|居左|居右)", text)
        if align_m:
            mapping = {"居中": "center", "左对齐": "left", "右对齐": "right",
                       "居左": "left", "居右": "right"}
            result["header"]["alignment"] = mapping.get(align_m.group(1), "center")

    # ── 页脚 ─────────────────────────────────────────────────────────────────
    footer_patterns = [
        r'页脚.*?' + _Q + r'(.*?)' + _QC,
        r'页脚.*?(?:内容|文字|写|加|设置|为)[：:]\s*(.+)',
    ]
    for pat in footer_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            result["footer"] = {"text": m.group(1).strip()}
            break

    # ── 页码 ─────────────────────────────────────────────────────────────────
    if re.search(r"页码", text):
        page_num_cfg: dict = {"position": "footer", "show_total": False}

        # 位置：页眉或页脚
        if re.search(r"页眉", text):
            page_num_cfg["position"] = "header"

        # 从第几页开始
        m_start = re.search(r"从第\s*(\d+)\s*页", text)
        if m_start:
            page_num_cfg["start_at"] = int(m_start.group(1))

        # 是否显示总页数
        if re.search(r"共\s*[X\d]*\s*页|总页|页数", text):
            page_num_cfg["show_total"] = True

        result["page_numbers"] = page_num_cfg

    # ── 目录 ─────────────────────────────────────────────────────────────────
    if re.search(r"目录", text) and re.search(r"增加|添加|插入|生成|加上|创建", text):
        toc_cfg: dict = {"title": "目录"}

        # 目录标题自定义
        m_title = re.search(r'目录.*?标题.*?[「『\u201c\u2018](.*?)[」』\u201d\u2019]', text)
        if m_title:
            toc_cfg["title"] = m_title.group(1).strip()

        # 插入位置
        m_pos = re.search(r"第\s*(\d+)\s*段", text)
        if m_pos:
            toc_cfg["insert_position"] = int(m_pos.group(1))

        result["toc"] = toc_cfg

    return result
