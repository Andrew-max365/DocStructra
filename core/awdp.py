from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any

import yaml
from docx import Document

AWDP_PROTOCOL = "AWDP-1.0"

_RE_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_RE_LEVEL1_HEADING = re.compile(r"^#\s+\S")
_RE_IMAGE = re.compile(r'!\[(.*?)\]\((\S+?)(?:\s+"([^"]+)")?\)')
_RE_HTML = re.compile(r"<[A-Za-z][^>]*>")
_RE_CODE_FENCE = re.compile(r"^```(.*)$")
_RE_ORDERED_LIST = re.compile(r"^\d+\.\s+")
_RE_UNORDERED_LIST = re.compile(r"^[-*+]\s+")
_RE_TABLE_SEPARATOR = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
_RE_ALLOWED_URL = re.compile(r"^https?://", re.IGNORECASE)
_RE_FRONT_MATTER = re.compile(r"^---\n(.*?)\n---(?:\n|$)", re.DOTALL)

# Inline formatting: order matters — bold+italic before bold, before italic
_RE_INLINE_FMT = re.compile(
    r"\*\*\*(.+?)\*\*\*"  # ***bold+italic***
    r"|\*\*(.+?)\*\*"     # **bold**
    r"|\*(.+?)\*"         # *italic*
    r"|~~(.+?)~~"         # ~~strikethrough~~
    r"|`(.+?)`"           # `inline code`
    r"|\[([^\]]+)\]\([^\)]+\)",  # [link text](url) — keep visible text only
)


@dataclass
class AWDPMarkdown:
    front_matter: dict[str, Any]
    body: str


class AWDPValidationError(ValueError):
    def __init__(self, errors: list[str]):
        self.errors = errors
        super().__init__("AWDP validation failed: " + " | ".join(errors))


def get_awdp_prompt_template() -> str:
    return (
        "请按照 AWDP-1.0 文档协议生成 Markdown 文档。\n\n"
        "要求：\n\n"
        "1) 文档必须以 YAML Front Matter 开头，并使用 --- 包裹。\n"
        "2) protocol 字段必须为 AWDP-1.0。\n"
        "3) 建议 Front Matter 字段：protocol、title、lang、author、date。\n"
        "4) Markdown 正文必须从一级标题开始。\n"
        "5) 只允许三层标题：#、##、###。\n"
        "6) 段落之间必须空一行。\n"
        "7) 表格必须使用标准 Markdown 表格（表头 + 分隔行）。\n"
        "8) 图片必须包含图标题，格式为 ![alt](url \"caption\")。\n"
        "9) 代码块必须使用 fenced code block，且必须声明语言，例如 ```python。\n"
        "10) 禁止使用 HTML 标签。\n"
    )


def parse_awdp_markdown(markdown_text: str) -> AWDPMarkdown:
    # Normalize Windows/mixed line endings so patterns always see \n
    text = (markdown_text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text.startswith("---\n"):
        raise AWDPValidationError(["缺少 YAML Front Matter（文档必须以 --- 开始）"])

    match = _RE_FRONT_MATTER.match(text)
    if not match:
        raise AWDPValidationError(["YAML Front Matter 未正确闭合（缺少结束 ---）"])
    yaml_text = match.group(1).strip()
    body = text[match.end():].lstrip("\n")

    if not yaml_text:
        raise AWDPValidationError(["YAML Front Matter 不能为空"])

    try:
        front_matter = yaml.safe_load(yaml_text)
    except Exception as e:
        raise AWDPValidationError([f"YAML Front Matter 解析失败: {e}"]) from e

    if not isinstance(front_matter, dict):
        raise AWDPValidationError(["YAML Front Matter 必须是键值对对象"])

    return AWDPMarkdown(front_matter=front_matter, body=body)


def validate_awdp_markdown(markdown_text: str) -> AWDPMarkdown:
    parsed = parse_awdp_markdown(markdown_text)
    errors: list[str] = []

    protocol = str(parsed.front_matter.get("protocol", "")).strip()
    if protocol != AWDP_PROTOCOL:
        errors.append(f"protocol 必须为 {AWDP_PROTOCOL}")

    body_lines = parsed.body.splitlines()
    first_non_empty = next((ln for ln in body_lines if ln.strip()), "")
    if not _RE_LEVEL1_HEADING.match(first_non_empty):
        errors.append("Markdown 正文必须从一级标题开始")

    in_code = False
    code_block_start_line = 0
    in_table = False
    for i, raw_line in enumerate(body_lines):
        line_no = i + 1
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        fence_match = _RE_CODE_FENCE.match(stripped)
        if fence_match:
            if not in_code:
                lang = fence_match.group(1).strip()
                if not lang:
                    errors.append(f"第 {line_no} 行代码块缺少语言声明")
                in_code = True
                code_block_start_line = line_no
            else:
                in_code = False
                code_block_start_line = 0
            continue

        if in_code:
            continue

        if _RE_HTML.search(line):
            errors.append(f"第 {line_no} 行包含 HTML 标签（协议禁止）")

        heading_match = _RE_HEADING.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            if level > 3:
                errors.append(f"第 {line_no} 行标题层级超过三级")
            in_table = False
            continue

        if not stripped:
            in_table = False
            continue

        if _RE_IMAGE.search(stripped):
            for m in _RE_IMAGE.finditer(stripped):
                caption = (m.group(3) or "").strip()
                if not caption:
                    errors.append(f"第 {line_no} 行图片缺少标题（caption）")
            in_table = False
            continue

        is_table_line = "|" in stripped
        if is_table_line:
            if not in_table:
                next_line = _next_line(body_lines, i)
                if not next_line or not _RE_TABLE_SEPARATOR.match(next_line):
                    errors.append(f"第 {line_no} 行表格不是标准 Markdown 表格（缺少分隔行）")
                in_table = True
            continue
        in_table = False

    if in_code:
        errors.append(f"代码块未闭合（起始于第 {code_block_start_line} 行，缺少结束 ```）")

    if errors:
        raise AWDPValidationError(errors)
    return parsed


def _apply_inline_formatting(paragraph: Any, text: str) -> None:
    """Add runs with bold/italic/code/strikethrough inline formatting to *paragraph*.

    Supported patterns:
      ***bold+italic*** | **bold** | *italic* | ~~strikethrough~~ | `code` | [link](url)
    """
    pos = 0
    for m in _RE_INLINE_FMT.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(1):  # ***bold+italic***
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # ~~strikethrough~~
            run = paragraph.add_run(m.group(4))
            run.font.strike = True
        elif m.group(5):  # `inline code`
            run = paragraph.add_run(m.group(5))
            run.font.name = "Courier New"
        elif m.group(6):  # [link text](url) — render visible text only
            paragraph.add_run(m.group(6))
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _flush_plain_paragraph(doc: Any, lines: list[str]) -> None:
    """Flush accumulated consecutive plain-text lines as a single paragraph."""
    if not lines:
        return
    text = " ".join(ln.strip() for ln in lines)
    text = text.strip()
    if text:
        p = doc.add_paragraph()
        _apply_inline_formatting(p, text)


def render_awdp_markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Validate and render AWDP-1.0 Markdown to a .docx byte string.

    Supports:
    - Headings (h1–h3)
    - Multi-line paragraphs (consecutive lines are merged into one paragraph)
    - Inline formatting: **bold**, *italic*, ***bold+italic***, ~~strike~~, `code`, [link](url)
    - Ordered and unordered lists
    - Block quotes
    - Fenced code blocks with language label
    - Tables (with grid style)
    - Images (rendered as placeholder + caption)
    """
    parsed = validate_awdp_markdown(markdown_text)
    doc = Document()

    body_lines = parsed.body.splitlines()
    i = 0
    para_buf: list[str] = []  # accumulates consecutive plain-text lines

    while i < len(body_lines):
        raw_line = body_lines[i]
        stripped = raw_line.strip()

        # Blank line: flush current paragraph buffer
        if not stripped:
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            i += 1
            continue

        # Fenced code block
        code_match = _RE_CODE_FENCE.match(stripped)
        if code_match:
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            lang = code_match.group(1).strip()
            i += 1
            code_lines: list[str] = []
            while i < len(body_lines):
                if _RE_CODE_FENCE.match(body_lines[i].strip()):
                    break
                code_lines.append(body_lines[i])
                i += 1
            title = doc.add_paragraph()
            title.add_run(f"代码（{lang}）").bold = True
            code_para = doc.add_paragraph("\n".join(code_lines))
            try:
                code_para.style = "No Spacing"
            except Exception:
                pass
            i += 1
            continue

        # Heading
        heading_match = _RE_HEADING.match(stripped)
        if heading_match:
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            i += 1
            continue

        # Table (header row followed immediately by separator row)
        next_line = _next_line(body_lines, i)
        if "|" in stripped and next_line and _RE_TABLE_SEPARATOR.match(next_line):
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            table_lines = [body_lines[i], body_lines[i + 1]]
            i += 2
            while i < len(body_lines):
                nxt = body_lines[i].strip()
                if not nxt or "|" not in nxt:
                    break
                table_lines.append(body_lines[i])
                i += 1
            rows = [_parse_table_row(x) for x in table_lines if "|" in x]
            if len(rows) >= 2:
                header = rows[0]
                data_rows = rows[2:] if len(rows) >= 3 else []
                col_count = max(1, len(header))
                table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                for c, cell_text in enumerate(header[:col_count]):
                    table.cell(0, c).text = cell_text
                for r, row_data in enumerate(data_rows, start=1):
                    for c in range(col_count):
                        table.cell(r, c).text = row_data[c] if c < len(row_data) else ""
            continue

        # Image placeholder
        image_match = _RE_IMAGE.search(stripped)
        if image_match:
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            alt = _sanitize_text(image_match.group(1))
            raw_url = (image_match.group(2) or "").strip()
            url = _sanitize_url(raw_url)
            caption = _sanitize_text(image_match.group(3) or "")
            img_p = doc.add_paragraph(f"【图片】{alt} ({url})")
            try:
                img_p.style = "Caption"
            except Exception:
                pass
            cap_p = doc.add_paragraph(f"图：{caption}")
            try:
                cap_p.style = "Caption"
            except Exception:
                pass
            i += 1
            continue

        # Ordered list item
        if _RE_ORDERED_LIST.match(stripped):
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            content = _RE_ORDERED_LIST.sub("", stripped, count=1).strip()
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, content)
            i += 1
            continue

        # Unordered list item
        if _RE_UNORDERED_LIST.match(stripped):
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            content = _RE_UNORDERED_LIST.sub("", stripped, count=1).strip()
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, content)
            i += 1
            continue

        # Block quote
        if stripped.startswith(">"):
            _flush_plain_paragraph(doc, para_buf)
            para_buf = []
            p = doc.add_paragraph(style="Intense Quote")
            _apply_inline_formatting(p, stripped.lstrip(">").strip())
            i += 1
            continue

        # Plain text — accumulate into paragraph buffer so consecutive lines
        # belonging to the same paragraph are merged into one.
        para_buf.append(raw_line)
        i += 1

    # Flush any remaining paragraph content
    _flush_plain_paragraph(doc, para_buf)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_table_row(line: str) -> list[str]:
    cells = line.strip()
    if cells.startswith("|"):
        cells = cells[1:]
    if cells.endswith("|"):
        cells = cells[:-1]
    return [c.strip() for c in cells.split("|")]


def _sanitize_text(value: str) -> str:
    text = (value or "").replace("\r", " ").replace("\n", " ").strip()
    return "".join(ch for ch in text if ch.isprintable())


def _sanitize_url(url: str) -> str:
    cleaned = _sanitize_text(url)
    if not _RE_ALLOWED_URL.match(cleaned):
        return "about:blank"
    return cleaned


def _next_line(lines: list[str], index: int) -> str:
    if index + 1 < len(lines):
        return lines[index + 1].strip()
    return ""
